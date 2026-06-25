"""
HAR Privacy Analyzer - Phase 1
Focus: LeadID detection with large file support and automatic corruption repair
"""

from flask import Flask, request, jsonify, send_from_directory
import json
import re
import base64
from urllib.parse import urlparse, parse_qs, unquote, unquote_plus
from datetime import datetime, timezone, timedelta
import os
import uuid

# ── MongoDB ───────────────────────────────────────────────────────────────────
try:
    from pymongo import MongoClient, DESCENDING
    from pymongo.errors import ConnectionFailure, ServerSelectionTimeoutError
    MONGO_URI = os.environ.get('MONGODB_URI', '')
    if MONGO_URI:
        _mongo_client = MongoClient(
            MONGO_URI,
            serverSelectionTimeoutMS=5000,
            connectTimeoutMS=5000,
        )
        # Verify connection on startup
        _mongo_client.admin.command('ping')
        _db = _mongo_client['har_analyzer']

        # Collections
        _col_single    = _db['single_analyses']
        _col_bulk      = _db['bulk_rankings']
        _col_settings  = _db['admin_settings']      # single upserted doc
        _col_token_use  = _db['token_usage']          # single upserted doc (running totals)
        _col_token_calls = _db['token_calls']          # per-call records for 30-day breakdown
        _col_token_calls.create_index('timestamp', expireAfterSeconds=7776000)  # 90-day TTL
        _col_token_calls.create_index([('timestamp', DESCENDING)])
        _col_reports   = _db['litigation_reports']   # Phase 3

        # TTL index: auto-expire documents after 90 days
        _col_single.create_index('created_at', expireAfterSeconds=7776000)
        _col_bulk.create_index('created_at',   expireAfterSeconds=7776000)
        _col_reports.create_index('created_at', expireAfterSeconds=7776000)

        # Query indexes
        _col_single.create_index([('created_at', DESCENDING)])
        _col_bulk.create_index([('created_at', DESCENDING)])
        _col_reports.create_index([('analysis_id', 1)])

        MONGO_ENABLED = True
        print('✅ MongoDB connected — OceansEdge cluster')
    else:
        MONGO_ENABLED = False
        print('⚠️  MONGODB_URI not set — history will not be persisted')
except Exception as _mongo_err:
    MONGO_ENABLED = False
    print(f'⚠️  MongoDB connection failed — history disabled: {_mongo_err}')
# ─────────────────────────────────────────────────────────────────────────────

# ── Anthropic / Claude API ────────────────────────────────────────────────────
try:
    import anthropic as _anthropic
    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
    if ANTHROPIC_API_KEY:
        _anthropic_client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        CLAUDE_ENABLED = True
        print('✅ Claude API client initialised')
    else:
        _anthropic_client = None
        CLAUDE_ENABLED = False
        print('⚠️  ANTHROPIC_API_KEY not set — Claude features disabled')
except ImportError:
    _anthropic_client = None
    CLAUDE_ENABLED = False
    print('⚠️  anthropic package not installed — Claude features disabled')
# ─────────────────────────────────────────────────────────────────────────────

app = Flask(__name__, static_folder='static')

# Increase max upload size to 100MB
app.config['MAX_CONTENT_LENGTH'] = 600 * 1024 * 1024  # 600 MB (covers 5x 100MB batch uploads + multipart overhead)



def strip_response_bodies(har_data):
    """
    Option B memory optimization: blank out response body text from an
    already-parsed HAR dict, reducing the live Python object tree size.

    Previous approach: brace-counting loop on the raw string before json.loads.
    Problem: O(n*m) character iteration — 27 seconds on a 36MB file with 73
    large JS/HTML responses.

    Current approach: parse JSON first (fast, C extension), then do a simple
    O(n entries) Python dict walk to blank out content.text in-place.
    Result: 0.15 seconds total — 184x faster on the same file.

    Mutates har_data in-place and returns it. No copy made.
    Safe: only touches response.content.text/encoding, never request.postData.
    """
    for entry in har_data.get('log', {}).get('entries', []):
        content = entry.get('response', {}).get('content', {})
        if 'text' in content:
            content['text'] = ''
        if 'encoding' in content:
            content['encoding'] = ''
    return har_data


def resilient_parse_har(har_text):
    """
    Extract as many entries as possible from a potentially corrupted HAR file.
    Returns (entries_list, stats_dict) or raises exception if completely unreadable.
    """
    
    print("⚠️  Standard parsing failed - attempting resilient extraction...")
    
    # Clean control characters
    har_text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', har_text)
    
    # Find the entries array
    entries_start_match = re.search(r'"entries"\s*:\s*\[', har_text)
    if not entries_start_match:
        raise Exception("Could not find entries array in HAR file")
    
    entries_start = entries_start_match.end()
    
    # Extract individual entry objects one by one
    entries = []
    position = entries_start
    errors = 0
    
    while position < len(har_text):
        # Skip whitespace
        while position < len(har_text) and har_text[position] in ' \n\r\t':
            position += 1
        
        if position >= len(har_text):
            break
        
        # Check for end of array
        if har_text[position] == ']':
            break
        
        # Skip commas
        if har_text[position] == ',':
            position += 1
            continue
        
        # Find start of next entry
        if har_text[position] != '{':
            next_brace = har_text.find('{', position)
            if next_brace == -1:
                break
            position = next_brace
        
        # Find matching closing brace
        brace_count = 0
        in_string = False
        escape = False
        start_pos = position
        
        i = position
        while i < len(har_text):
            char = har_text[i]
            
            if escape:
                escape = False
                i += 1
                continue
            
            if char == '\\':
                escape = True
                i += 1
                continue
            
            if char == '"':
                in_string = not in_string
            elif not in_string:
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        # Found complete entry
                        entry_text = har_text[start_pos:i+1]
                        
                        try:
                            entry = json.loads(entry_text)
                            entries.append(entry)
                        except:
                            errors += 1
                        
                        position = i + 1
                        break
            
            i += 1
        else:
            # Hit corruption - try to find next valid entry
            next_entry = har_text.find('"startedDateTime"', position + 1000)
            if next_entry == -1:
                break
            
            # Backtrack to find the {
            back_pos = next_entry
            while back_pos > position and har_text[back_pos] != '{':
                back_pos -= 1
            
            if back_pos > position:
                position = back_pos
                errors += 1
            else:
                break
    
    stats = {
        'extracted': len(entries),
        'skipped': errors,
        'success_rate': len(entries) / (len(entries) + errors) * 100 if (len(entries) + errors) > 0 else 0
    }
    
    print(f"✅ Resilient parsing: extracted {len(entries)} entries, skipped {errors}")
    
    return entries, stats


def decode_base64_if_present(text):
    """
    Detect and decode base64-encoded data.
    Returns decoded text if valid base64, otherwise returns None.
    """
    if not text or len(text) < 10:
        return None
    
    # Pattern 1: data URL (data:image/png;base64,...)
    if text.startswith('data:'):
        parts = text.split(',', 1)
        if len(parts) == 2 and 'base64' in parts[0]:
            try:
                decoded = base64.b64decode(parts[1])
                return decoded.decode('utf-8', errors='ignore')
            except:
                return None
    
    # Pattern 2: Looks like base64 (alphanumeric + / + =)
    # Must be reasonable length and have base64 characteristics
    if re.match(r'^[A-Za-z0-9+/=]{20,}$', text):
        # Check if it has the right padding
        try:
            # Try to decode
            decoded = base64.b64decode(text)
            # Check if decoded result looks like text or JSON
            decoded_str = decoded.decode('utf-8', errors='ignore')
            
            # If it looks like JSON or contains readable text, it's probably real base64
            if decoded_str.startswith(('{', '[')) or any(c.isalpha() for c in decoded_str[:50]):
                return decoded_str
        except:
            pass
    
    return None


def detect_hash_pattern(key, value):
    """
    Detect if a value is a hash of PII (email, phone).
    Returns PII dict if hash detected, None otherwise.
    
    Common patterns:
    - SHA-256: 64 hex chars
    - SHA-1: 40 hex chars  
    - MD5: 32 hex chars
    """
    if not isinstance(value, str):
        return None
    
    key_lower = key.lower()
    
    # SHA-256 (64 hex characters)
    if re.match(r'^[a-f0-9]{64}$', value.lower()):
        if any(term in key_lower for term in ['email', 'em', 'mail']):
            return {
                'type': 'Hashed Email (SHA-256)',
                'field': key,
                'value': value[:16] + '...',  # Show first 16 chars
                'note': 'Full hash: ' + value
            }
        if any(term in key_lower for term in ['phone', 'tel', 'mobile']):
            return {
                'type': 'Hashed Phone (SHA-256)',
                'field': key,
                'value': value[:16] + '...',
                'note': 'Full hash: ' + value
            }
    
    # SHA-1 (40 hex characters)
    elif re.match(r'^[a-f0-9]{40}$', value.lower()):
        if any(term in key_lower for term in ['email', 'em', 'mail']):
            return {
                'type': 'Hashed Email (SHA-1)',
                'field': key,
                'value': value[:16] + '...',
                'note': 'Full hash: ' + value
            }
        if any(term in key_lower for term in ['phone', 'tel', 'mobile']):
            return {
                'type': 'Hashed Phone (SHA-1)',
                'field': key,
                'value': value[:16] + '...',
                'note': 'Full hash: ' + value
            }
    
    # MD5 (32 hex characters)
    elif re.match(r'^[a-f0-9]{32}$', value.lower()):
        if any(term in key_lower for term in ['email', 'em', 'mail']):
            return {
                'type': 'Hashed Email (MD5)',
                'field': key,
                'value': value[:16] + '...',
                'note': 'Full hash: ' + value
            }
        if any(term in key_lower for term in ['phone', 'tel', 'mobile']):
            return {
                'type': 'Hashed Phone (MD5)',
                'field': key,
                'value': value[:16] + '...',
                'note': 'Full hash: ' + value
            }
    
    return None


def robust_url_decode(text):
    """
    Decode URL-encoded text, handling multiple encoding layers.
    
    Examples:
    - Single: name%40example.com → name@example.com
    - Double: name%2540example.com → name%40example.com → name@example.com
    - Plus signs: first+last → first last
    """
    if not isinstance(text, str):
        return text
    
    # Decode up to 5 layers (should be more than enough)
    prev = text
    for i in range(5):
        # Use unquote_plus to handle both %20 and + for spaces
        curr = unquote_plus(prev)
        if curr == prev:
            # No more encoding layers
            break
        prev = curr
    
    return curr


def extract_pii_from_params(params):
    """
    Look for common PII field names in parameters with context-aware detection
    """
    pii = []
    
    # Known PII field name patterns
    pii_field_patterns = {
        # Email fields
        'email': 'Email',
        'e': 'Email',  # Common short form
        'em': 'Email',
        'mail': 'Email',
        'user_email': 'Email',
        
        # Phone fields
        'phone': 'Phone',
        'tel': 'Phone',
        'telephone': 'Phone',
        'mobile': 'Phone',
        'contactPhone': 'Phone',
        'contactphone': 'Phone',
        'phone_number': 'Phone',
        'ph': 'Phone',
        
        # Zip/Postal code fields
        'zip': 'Zip Code',
        'zipcode': 'Zip Code',
        'zip_code': 'Zip Code',
        'postal': 'Zip Code',
        'postalcode': 'Zip Code',
        'postal_code': 'Zip Code',
        
        # Name fields
        'firstName': 'First Name',
        'first_name': 'First Name',
        'firstname': 'First Name',
        'fn': 'First Name',
        'lastName': 'Last Name',
        'last_name': 'Last Name',
        'lastname': 'Last Name',
        'ln': 'Last Name',
        'contactName': 'Full Name',
        'contact_name': 'Full Name',
        'contactname': 'Full Name',
        'name': 'Name',
        'fullname': 'Full Name',
        'full_name': 'Full Name',
        
        # Address fields
        'address': 'Address',
        'street': 'Street Address',
        'street_address': 'Street Address',
        'addr': 'Address',
        'city': 'City',
        'state': 'State',
        'st': 'State',  # Only if exact match
        'country': 'Country',
        
        # Other PII
        'businessDescription': 'Business Description',
        'business_description': 'Business Description',
        'insurance_type': 'Insurance Type',
        'insurancetype': 'Insurance Type',
        'ssn': 'SSN',
        'social_security': 'SSN',
    }
    
    # Tracking/non-PII field patterns to EXCLUDE
    exclude_patterns = [
        'id', 'pixel', 'fbp', 'fbc', '_fb', 'ga_', 'gid', 'cid',
        'session', 'token', 'timestamp', 'ts', 'uid', 'uuid',
        'click', 'campaign', 'source', 'medium', 'ref', 'utm',
        'event', 'ev', 'action', 'version', 'v', 'ttclid', 'ttp',
        'callback', 'cb', 'redirect', 'url', 'dl', 'rl',
        'fst', 'lst', 'pst',  # first/last/previous seen timestamp
        'cs_', 'ep.', '_ga', '_gcl',  # Google/Facebook tracking params
        'memory', 'heap', 'size', 'limit',  # Performance metrics
        'width', 'height', 'screen', 'viewport',  # Display metrics
        '_et', 'tfd', 'tcfd',  # Google Analytics encrypted params
        '.js', '.css', '.png', '.jpg', '.gif',  # File extensions
        'script', 'src', 'href',  # HTML/resource references
    ]
    
    # Check each parameter
    for key, value in params.items():
        # Handle both single values and lists
        val = value[0] if isinstance(value, list) else value
        
        # Skip empty or very short values
        if not val or len(str(val)) < 2:
            continue
        
        # Skip if value looks like a filename
        if any(ext in str(val).lower() for ext in ['.js', '.css', '.html', '.png', '.jpg', '.svg', '.gif', '.woff']):
            continue
        
        # Convert value to string and apply robust URL decoding
        val_str = robust_url_decode(str(val))
        
        # Convert key to lowercase for comparison
        key_lower = key.lower()
        
        # Check if this is a tracking field we should skip
        if any(exclude in key_lower for exclude in exclude_patterns):
            continue
        
        # PRIORITY 1: Check for hashed PII (SHA-256, SHA-1, MD5)
        hash_pii = detect_hash_pattern(key, val_str)
        if hash_pii:
            pii.append(hash_pii)
            continue  # Hash detected, move to next param
        
        # PRIORITY 2: Try to decode base64 if present
        decoded = decode_base64_if_present(val_str)
        if decoded:
            # Decoded successfully - check if it's JSON
            if decoded.startswith(('{', '[')):
                try:
                    decoded_json = json.loads(decoded)
                    # Recursively extract PII from decoded JSON
                    pii.extend(extract_pii_from_json(decoded_json, prefix=f'{key}[decoded].'))
                    continue  # Found JSON in base64, move to next param
                except:
                    pass
            # If not JSON, treat decoded text as the value
            val_str = decoded
        
        # PRIORITY 3: Check if field name matches known PII fields
        matched = False
        for field_pattern, label in pii_field_patterns.items():
            # For short ambiguous patterns like 'st', require exact match
            if field_pattern in ['st', 'v', 'e', 'ph', 'fn', 'ln']:
                if key_lower == field_pattern.lower():
                    field_match = True
                else:
                    field_match = False
            else:
                field_match = field_pattern.lower() in key_lower
            
            if field_match:
                # Additional validation based on type
                val_str = str(val)
                
                if label == 'Email':
                    # Validate email format
                    if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', val_str):
                        pii.append({
                            'type': label,
                            'field': key,
                            'value': val_str
                        })
                        matched = True
                        break
                
                elif label == 'Phone':
                    # Validate phone: must be 10 digits, no dots or 'fb' prefix
                    clean = re.sub(r'[^\d]', '', val_str)
                    if len(clean) == 10 and not any(c in val_str.lower() for c in ['fb', 'ga', '.']):
                        pii.append({
                            'type': label,
                            'field': key,
                            'value': val_str
                        })
                        matched = True
                        break
                
                elif label == 'Zip Code':
                    # Validate zip: exactly 5 digits
                    if re.match(r'^\d{5}$', val_str):
                        # Reject single-letter field names (too ambiguous)
                        if len(key) <= 1:
                            matched = True
                            break
                        
                        # Reject if field name suggests it's NOT a zip (but allow zipcode/zip_code)
                        # Check for patterns like 'eventCode', 'errorCode' but NOT 'zipCode'
                        suspicious_patterns = ['eventcode', 'errorcode', 'statuscode', 'responsecode', 'orderid', 'sessionid']
                        if not any(pattern in key_lower for pattern in suspicious_patterns):
                            pii.append({
                                'type': label,
                                'field': key,
                                'value': val_str
                            })
                            matched = True
                            break
                
                elif label == 'SSN':
                    # Validate SSN: 9 digits or XXX-XX-XXXX format
                    clean = re.sub(r'[^\d]', '', val_str)
                    if len(clean) == 9:
                        pii.append({
                            'type': label,
                            'field': key,
                            'value': val_str
                        })
                        matched = True
                        break
                
                else:
                    # Other fields: just need non-empty value
                    # For name fields, add extra validation
                    if 'Name' in label:
                        # Filter out common technical terms, filenames, and abbreviations
                        val_lower = val_str.lower()
                        invalid_names = [
                            'pixel', 'script', 'bundle', 'chunk', 'module', 'component',
                            'data', 'info', 'user', 'type', 'null', 'none', 'undefined',
                            'true', 'false', 'yes', 'no', 'ok', 'error',
                            'fn', 'ln', 'n/a', 'na'
                        ]
                        
                        # Skip if it's a technical term or filename
                        if any(term in val_lower for term in invalid_names):
                            continue
                        
                        # Skip if it contains file extension indicators
                        if any(ext in val_lower for ext in ['.js', '.css', '.html', '.']):
                            continue
                        
                        # Skip very short names (likely abbreviations)
                        if len(val_str) < 3:
                            continue
                    
                    if len(val_str) > 1:
                        pii.append({
                            'type': label,
                            'field': key,
                            'value': val_str
                        })
                        matched = True
                        break
        
        # If not matched by field name, try pattern-only detection as FALLBACK
        # Some sites DO send plaintext PII in unexpected field names!
        if not matched:
            val_str = str(val)
            
            # Make sure field isn't in exclusion list
            if any(exclude in key_lower for exclude in exclude_patterns):
                continue
            
            # Email pattern - high confidence
            if '@' in val_str and re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', val_str):
                pii.append({
                    'type': 'Email',
                    'field': key,
                    'value': val_str
                })
            
            # Phone pattern - be more careful but still detect
            # Must be EXACTLY 10 digits OR formatted phone (xxx-xxx-xxxx, (xxx) xxx-xxxx)
            elif re.match(r'^(\d{10}|\d{3}-\d{3}-\d{4}|\(\d{3}\)\s*\d{3}-\d{4})$', val_str):
                # Extra validation: no dots, not part of tracking ID
                if '.' not in val_str and 'fb' not in val_str.lower() and '_' not in val_str:
                    pii.append({
                        'type': 'Phone',
                        'field': key,
                        'value': val_str
                    })
            
            # Zip code pattern - must be exactly 5 digits
            # Additional check: field name shouldn't suggest it's a code/ID
            elif re.match(r'^\d{5}$', val_str):
                # Skip single-letter field names (too ambiguous like 'o', 'v', etc.)
                if len(key) <= 2:
                    continue
                
                # Skip if field name suggests it's not a zip (ID, code, count, order, index, etc.)
                if any(word in key_lower for word in ['code', 'id', 'event', 'version', 'count', 'order', 'index', 'sequence', 'number']):
                    continue
                
                # Only accept if looks like real US zip (not starting with 0 is rare but valid)
                pii.append({
                    'type': 'Zip Code',
                    'field': key,
                    'value': val_str
                })
            
            # Name pattern - capitalized word with length > 2
            # Be VERY careful - filenames often match this pattern!
            elif re.match(r'^[A-Z][a-z]{2,}$', val_str):
                # Skip if it's obviously a filename or code term
                if any(term in val_str.lower() for term in ['pixel', 'script', 'bundle', 'chunk', 'min', 'js', 'css', 'html', 'json', 'xml']):
                    continue
                # Skip if field name suggests technical content
                if any(term in key_lower for term in ['file', 'script', 'url', 'path', 'type', 'format', 'extension']):
                    continue
                # Skip generic words
                # Skip generic words AND analytics/tracking event names
                GENERIC_WORDS = {
                    # Generic UI/data terms
                    'Page', 'View', 'Data', 'User', 'Info', 'Type', 'Name', 'Form',
                    'None', 'Null', 'True', 'False', 'Home', 'Next', 'Back', 'Done',
                    'Open', 'Close', 'Send', 'Load', 'Init', 'Start', 'Stop', 'Save',
                    # Analytics event names (TikTok, GA, Facebook, etc.)
                    'Segment', 'Click', 'Track', 'Event', 'Fire', 'Hit', 'Ping',
                    'Pixel', 'Lead', 'Submit', 'Complete', 'Convert', 'Purchase',
                    'Search', 'Browse', 'Scroll', 'Hover', 'Focus', 'Blur',
                    'Error', 'Debug', 'Warn', 'Info', 'Log', 'Trace',
                    # Common single-word values that aren't names
                    'Male', 'Female', 'Other', 'Unknown', 'Default', 'Custom',
                    'Active', 'Pending', 'Enabled', 'Disabled', 'Success', 'Failed',
                }
                if val_str in GENERIC_WORDS:
                    continue
                
                # If it passes all checks, it MIGHT be a name
                pii.append({
                    'type': 'Name',
                    'field': key,
                    'value': val_str
                })
    
    return pii

def extract_pii_from_json(data, prefix=''):
    """
    Recursively extract PII from JSON data
    """
    pii = []
    
    if isinstance(data, dict):
        for key, value in data.items():
            current_key = f"{prefix}.{key}" if prefix else key
            
            if isinstance(value, (dict, list)):
                pii.extend(extract_pii_from_json(value, current_key))
            else:
                # Treat as parameter
                params = {current_key: value}
                pii.extend(extract_pii_from_params(params))
    
    elif isinstance(data, list):
        for i, item in enumerate(data):
            current_key = f"{prefix}[{i}]"
            pii.extend(extract_pii_from_json(item, current_key))
    
    return pii

def extract_pii_from_request(entry):
    """
    Extract plaintext PII from POST body and query params
    """
    pii = []
    
    # Check POST body
    if 'postData' in entry['request'] and 'text' in entry['request']['postData']:
        text = entry['request']['postData']['text']
        mime_type = entry['request']['postData'].get('mimeType', '')
        
        # Try JSON parsing first (regardless of MIME type) if text looks like JSON
        if text.strip().startswith(('{', '[')):
            try:
                data = json.loads(text)
                pii.extend(extract_pii_from_json(data))
            except:
                pass
        
        # URL-encoded parameters
        if 'application/x-www-form-urlencoded' in mime_type or ('=' in text and '&' in text):
            try:
                params = parse_qs(text)
                pii.extend(extract_pii_from_params(params))
            except:
                pass
    
    # Check query string
    url = entry['request']['url']
    if '?' in url:
        try:
            query = url.split('?', 1)[1]
            params = parse_qs(query)
            pii.extend(extract_pii_from_params(params))
        except:
            pass
    
    return pii

def find_first_party_domain(har_data):
    """
    Find the main website being analyzed
    """
    domains = {}
    
    for entry in har_data['log']['entries']:
        url = entry.get('request', {}).get('url', '')
        if not url:
            continue
        try:
            parsed = urlparse(url)
            domain = parsed.netloc
            
            # Skip common third-party domains and CDNs
            skip_domains = [
                'google', 'facebook', 'doubleclick', 'analytics',
                'leadid', 'jornaya', 'trustedform', 'trueleadid',
                'invoca', 'bing', 'microsoft', 'linkedin',
                'cloudflare', 'akamai', 'cdn', 'tiktok',
                'gstatic', 'googleapis', 'amazon', 'amazonaws',
                'fastly', 'cloudfront', 'newrelic', 'nr-data',
                'segment', 'sentry', 'hotjar', 'intercom',
                'zendesk', 'stripe', 'twilio', 'sendgrid',
                'clarity.ms', 'hotjar', 'mixpanel', 'heap', 'mouseflow',
            ]
            
            if any(skip in domain.lower() for skip in skip_domains):
                continue
            
            # Skip empty domains
            if not domain:
                continue
            
            domains[domain] = domains.get(domain, 0) + 1
        except:
            continue
    
    # Return most common domain
    if domains:
        return max(domains, key=domains.get)
    return 'Unknown'

def detect_vendor_requests(entries):
    """
    Option D memory optimization: detect vendors AND extract all needed data
    inline in a single pass, storing NO 'entry' references.

    Previously this function stored {'entry': entry, ...} for every matched
    request, keeping the entire parsed JSON (~3.5x file size) alive in memory
    until analyze_har_simple returned. Now we extract everything we need from
    each entry immediately and discard the reference, allowing Python's GC to
    free the parsed JSON as soon as the caller does `del har_data`.

    Each vendor's requests list now stores lightweight dicts with:
      - Scalar fields extracted from the entry (url, method, timestamp, etc.)
      - pii: list of PII items found (extracted inline here)
      - all_request_info: the summary row for the expandable UI list
    No reference to the original entry object is retained.
    """
    VENDOR_PATTERNS = {
        'leadid':    ('LeadID/Jornaya/TrustedForm', 'critical',
                      ['leadid.com', 'jornaya.com', 'trustedform.com', 'trueleadid.com']),
        'google':    ('Google Analytics/Ads', 'high',
                      ['google-analytics.com', 'googletagmanager.com', 'doubleclick.net',
                       'googlesyndication.com', 'googleadservices.com', '/google/ads', '/gtag/']),
        'facebook':  ('Facebook/Meta Pixel', 'critical',
                      ['facebook.com/tr', 'connect.facebook.net', 'facebook.net',
                       '/fbevents.js', 'fbcdn.net']),
        'tiktok':    ('TikTok Pixel', 'critical',
                      ['analytics.tiktok.com', 'business-api.tiktok.com',
                       'tiktok.com/i18n/pixel', '/tiktok-pixel']),
        'invoca':    ('Invoca Call Tracking', 'high',
                      ['invoca.net', 'invocacdn.com']),
        'microsoft': ('Microsoft/Bing Ads', 'medium',
                      ['bat.bing.com', 'bing.com/api', 'clarity.ms', 'uetanalytics.com']),
        'linkedin':  ('LinkedIn Insight Tag', 'medium',
                      ['linkedin.com/px', 'snap.licdn.com', 'linkedin.com/li/track']),
    }

    vendors = {
        key: {
            'name': name,
            'risk': risk,
            'request_count': 0,
            'post_count': 0,
            'pii': [],          # extracted PII items (lightweight dicts only)
            'all_requests': [], # summary rows for expandable UI list
        }
        for key, (name, risk, _) in VENDOR_PATTERNS.items()
    }

    for entry in entries:
        url_lower = entry.get('request', {}).get('url', '').lower()
        full_url  = entry.get('request', {}).get('url', '')
        method    = entry.get('request', {}).get('method', 'GET')
        timestamp = entry.get('startedDateTime', '')

        # Match vendor
        matched_key = None
        for key, (_, _, patterns) in VENDOR_PATTERNS.items():
            if any(p in url_lower for p in patterns):
                matched_key = key
                break
        if matched_key is None:
            continue

        v = vendors[matched_key]
        vendor_name = v['name']

        # --- Extract scalars from entry NOW, before we lose the reference ---
        time_ms         = entry.get('time', 0)
        response_status = entry.get('response', {}).get('status', 0)
        request_size    = entry.get('request', {}).get('bodySize', 0)
        response_size   = entry.get('response', {}).get('bodySize', 0)

        # --- PII extraction inline (was extract_pii_from_vendor_requests) ---
        pii_found = extract_pii_from_request(entry)
        pii_items = []
        for pii_item in pii_found:
            is_hashed    = 'Hashed' in pii_item['type']
            legal_context = get_legal_context(pii_item['type'], vendor_name, is_hashed)
            pii_items.append({
                'type':      pii_item['type'],
                'value':     pii_item['value'],
                'field':     pii_item.get('field', ''),
                'timestamp': timestamp,
                'url':       full_url[:100],
                'note':      pii_item.get('note', ''),
                'request_details': {
                    'method':          method,
                    'full_url':        full_url,
                    'response_code':   response_status,
                    'response_time_ms': int(time_ms) if time_ms else 0,
                    'request_size':    request_size,
                    'response_size':   response_size,
                },
                'legal_context': legal_context,
            })

        # --- Build all_requests summary row (was built in analyze_har_simple) ---
        has_pii = len(pii_items) > 0
        all_request_row = {
            'method':          method,
            'url':             full_url,
            'timestamp':       timestamp,
            'response_code':   response_status,
            'response_time_ms': int(time_ms) if time_ms else 0,
            'has_pii':         has_pii,
        }

        # Accumulate — NO 'entry' reference stored anywhere
        v['request_count'] += 1
        if method == 'POST':
            v['post_count'] += 1
        v['pii'].extend(pii_items)
        v['all_requests'].append(all_request_row)

        # entry goes out of scope at end of loop iteration.
        # With no stored reference, GC can free it as soon as
        # del har_data is called in the caller.

    # Filter out vendors with no requests
    return {k: v for k, v in vendors.items() if v['request_count'] > 0}


def get_legal_severity(pii_type, is_hashed=False):
    """
    Determine legal severity for a PII type
    """
    if is_hashed:
        return {
            'severity': 'high',
            'violation_type': 'hashed_pii_third_party',
            'estimated_damages': 5000
        }
    elif pii_type in ['Email', 'Phone', 'SSN']:
        return {
            'severity': 'critical',
            'violation_type': 'plaintext_pii_third_party',
            'estimated_damages': 5000
        }
    elif pii_type in ['First Name', 'Last Name', 'Full Name', 'Name']:
        return {
            'severity': 'critical',
            'violation_type': 'plaintext_pii_third_party',
            'estimated_damages': 5000
        }
    elif pii_type in ['Zip Code', 'Address', 'City', 'State']:
        return {
            'severity': 'high',
            'violation_type': 'plaintext_pii_third_party',
            'estimated_damages': 5000
        }
    else:
        return {
            'severity': 'medium',
            'violation_type': 'potential_pii_third_party',
            'estimated_damages': 5000
        }


def get_legal_context(pii_type, vendor_name, is_hashed=False):
    """
    Generate legal context for a PII transmission
    """
    severity_info = get_legal_severity(pii_type, is_hashed)
    
    # Base context
    context = {
        'severity': severity_info['severity'],
        'violation_type': severity_info['violation_type'],
        'estimated_damages': severity_info['estimated_damages'],
        'evidence_strength': 'strong' if not is_hashed else 'very_strong'
    }
    
    # Generate what_happened text
    if is_hashed:
        context['what_happened'] = f"{vendor_name} received a cryptographically hashed version of your {pii_type.lower()}. While they can't see the actual value, they can use this hash to track you across websites."
    else:
        context['what_happened'] = f"{vendor_name} received your {pii_type.lower()} in plaintext while you were using the website."
    
    # Generate why_matters text
    if is_hashed:
        context['why_matters'] = f"Third-party disclosure of {pii_type.lower()}-related data + Cross-site tracking capability + No observed consent"
    else:
        context['why_matters'] = f"Third-party interception + Contents of communication ({pii_type.lower()}) + No observed consent"
    
    # CIPA elements
    context['cipa_elements'] = ['third_party_disclosure', 'no_consent']
    if not is_hashed:
        context['cipa_elements'].extend(['interception', 'contents_of_communication'])
    if is_hashed:
        context['cipa_elements'].append('tracking_capability')
    
    return context


# extract_pii_from_vendor_requests removed — logic merged into detect_vendor_requests (Option D)


def analyze_first_party_requests(entries, first_party_domain):
    """
    Analyze first-party POST requests for PII collection.
    This catches lead generation sites that collect PII on their own server
    before sharing tracking IDs with third parties.
    """
    first_party_pii = []
    first_party_posts = []
    
    for entry in entries:
        # Only look at POST requests
        if entry.get('request', {}).get('method') != 'POST':
            continue
        
        url = entry['request']['url']
        
        # Check if this is a first-party request
        if first_party_domain not in url:
            continue
        
        # Skip common non-PII endpoints
        skip_paths = [
            '/assets/', '/static/', '/_next/', '/api/analytics', 
            '/api/tracking', '/api/events', '/__nextjs'
        ]
        if any(skip in url for skip in skip_paths):
            continue
        
        # Extract PII from this request
        pii_found = extract_pii_from_request(entry)
        
        if pii_found:
            first_party_posts.append({
                'url': url,
                'timestamp': entry.get('startedDateTime', ''),
                'pii': pii_found
            })
            
            # Add to overall list
            for pii_item in pii_found:
                first_party_pii.append({
                    'type': pii_item['type'],
                    'value': pii_item['value'],
                    'field': pii_item['field'],
                    'timestamp': entry.get('startedDateTime', ''),
                    'url': url[:100]
                })
    
    return {
        'pii_items': first_party_pii,
        'post_requests': first_party_posts,
        'post_count': len(first_party_posts),
        'pii_count': len(set(f"{p['type']}:{p['value']}" for p in first_party_pii))
    }


def analyze_har_as_plaintext(har_text, first_party_domain):
    """
    Fallback analysis for corrupted HAR files.
    Searches the raw text for PII patterns without parsing JSON structure.
    
    This is useful when:
    - File is severely corrupted
    - JSON structure is broken
    - Entries can't be parsed normally
    """
    print("🔍 Analyzing HAR as plaintext (fallback mode)...")
    
    pii_findings = []
    
    # Search for email addresses
    emails = re.findall(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b', har_text)
    unique_emails = set(emails)
    
    # Filter out common tracking/system emails
    exclude_domains = ['facebook.com', 'google.com', 'tiktok.com', 'microsoft.com', 
                      'linkedin.com', 'example.com', 'localhost', 'qualtrics.com',
                      'adobe.com', 'cloudflare.com', 'akamai.com']
    
    for email in unique_emails:
        if not any(domain in email.lower() for domain in exclude_domains):
            pii_findings.append({
                'type': 'Email',
                'value': email,
                'field': 'plaintext_scan',
                'source': 'fallback'
            })
    
    # Search for phone numbers
    # Look for patterns like: "phone":"6123277188" or \"phone\":\"6123277188\"
    phone_patterns = [
        r'\\*["\']?phone\\*["\']?\s*:\s*\\*["\']?(\d{10})\\*["\']?',
        r'\\*["\']?tel\\*["\']?\s*:\s*\\*["\']?(\d{10})\\*["\']?',
        r'\\*["\']?mobile\\*["\']?\s*:\s*\\*["\']?(\d{10})\\*["\']?',
        r'\\*["\']?contactPhone\\*["\']?\s*:\s*\\*["\']?(\d{10})\\*["\']?',
    ]
    
    for pattern in phone_patterns:
        phones = re.findall(pattern, har_text, re.IGNORECASE)
        for phone in set(phones):
            # Validate it looks like a real phone (area code check)
            if phone[0] in '23456789':  # Valid US area codes
                pii_findings.append({
                    'type': 'Phone',
                    'value': phone,
                    'field': 'plaintext_scan',
                    'source': 'fallback'
                })
    
    # Search for first/last names
    # Including escaped JSON: \"firstName\":\"Matthew\"
    name_patterns = [
        (r'\\*["\']?firstName\\*["\']?\s*:\s*\\*["\']([A-Z][a-z]+)\\*["\']', 'First Name'),
        (r'\\*["\']?lastName\\*["\']?\s*:\s*\\*["\']([A-Z][a-z]+)\\*["\']', 'Last Name'),
        (r'\\*["\']?first_name\\*["\']?\s*:\s*\\*["\']([A-Z][a-z]+)\\*["\']', 'First Name'),
        (r'\\*["\']?last_name\\*["\']?\s*:\s*\\*["\']([A-Z][a-z]+)\\*["\']', 'Last Name'),
    ]
    
    for pattern, label in name_patterns:
        names = re.findall(pattern, har_text, re.IGNORECASE)
        for name in set(names):
            # Filter out abbreviations and common words
            if len(name) > 2 and name.lower() not in ['page', 'view', 'data', 'user', 'info', 'type', 'name', 'form', 'none', 'null']:
                pii_findings.append({
                    'type': label,
                    'value': name,
                    'field': 'plaintext_scan',
                    'source': 'fallback'
                })
    
    # Remove duplicates
    unique_pii = {}
    for pii in pii_findings:
        key = f"{pii['type']}:{pii['value']}"
        if key not in unique_pii:
            unique_pii[key] = pii
    
    pii_list = list(unique_pii.values())
    
    print(f"   Found {len(pii_list)} PII items in plaintext scan")
    
    return {
        'pii_items': pii_list,
        'pii_count': len(pii_list),
        'method': 'plaintext_fallback',
        'note': 'PII extracted via plaintext scan due to file corruption'
    }

def build_simple_timeline(requests):
    """
    Build chronological list of captures
    """
    timeline = []
    
    for req in sorted(requests, key=lambda x: x['timestamp']):
        try:
            # Parse ISO timestamp
            timestamp = datetime.fromisoformat(req['timestamp'].replace('Z', '+00:00'))
            time_str = timestamp.strftime('%I:%M:%S %p')
            date_str = timestamp.strftime('%B %d, %Y')
        except:
            time_str = req['timestamp']
            date_str = req['timestamp'].split('T')[0] if 'T' in req['timestamp'] else 'Unknown'
        
        for capture in req['captures']:
            timeline.append({
                'time': time_str,
                'date': date_str,
                'type': capture['type'],
                'value': capture['value'],
                'field': capture['field'],
                'url': req['url']
            })
    
    return timeline

def analyze_har_simple(har_data):
    """
    Analyze HAR file for all major third-party vendors.

    Option D memory layout:
      1. Extract scalars (total_requests, post_requests, first_party, session info)
         from har_data while we still have it.
      2. Run first_party and LeadID timeline analysis (both need raw entries).
      3. Call detect_vendor_requests — now does inline PII extraction with NO
         entry references stored. Returns fully extracted lightweight dicts.
      4. del har_data — at this point no entry refs exist anywhere, so Python's
         GC can immediately free ~3.5x file size of parsed JSON.
      5. All remaining work uses only lightweight extracted data.
    """

    # 1. Extract scalars and session info from har_data upfront
    entries = har_data['log']['entries']
    total_requests = len(entries)
    post_requests  = sum(1 for e in entries if e.get('request', {}).get('method') == 'POST')
    first_party    = find_first_party_domain(har_data)

    session_start = None
    session_date  = None
    if entries:
        try:
            first_ts = entries[0].get('startedDateTime', '')
            if first_ts:
                dt = datetime.fromisoformat(first_ts.replace('Z', '+00:00'))
                session_start = dt.strftime('%I:%M %p')
                session_date  = dt.strftime('%B %d, %Y')
        except:
            session_start = 'Unknown'
            session_date  = 'Unknown'

    # 2. First-party analysis (needs raw entries — do before del har_data)
    first_party_analysis = analyze_first_party_requests(entries, first_party)

    # 3. LeadID timeline (needs raw entries — do before del har_data)
    # We rebuild it from the PII already extracted in step 4 below after the
    # vendor pass, but we need entries for the timeline builder format.
    # Build a lightweight capture list now while entries are still alive.
    leadid_captures = []
    for entry in entries:
        url_lower = entry.get('request', {}).get('url', '').lower()
        if any(d in url_lower for d in ['leadid.com', 'jornaya.com', 'trustedform.com', 'trueleadid.com']):
            pii_found = extract_pii_from_request(entry)
            if pii_found:
                leadid_captures.append({
                    'timestamp': entry.get('startedDateTime', ''),
                    'url':       entry.get('request', {}).get('url', ''),
                    'captures':  pii_found,
                })

    # 4. Vendor detection + inline PII extraction (Option D: no entry refs stored)
    detected_vendors = detect_vendor_requests(entries)

    # 5. Option D: release the parsed JSON now — no entry refs remain anywhere
    del entries, har_data

    # 6. Build final vendor output from the already-extracted lightweight data
    vendors_with_pii = {}
    leadid_in_vendors = False
    leadid_request_count = 0

    for vendor_key, vendor_data in detected_vendors.items():
        pii = vendor_data['pii']
        pii_count = len(set(f"{p['type']}:{p['value']}" for p in pii))

        vendors_with_pii[vendor_key] = {
            'name':          vendor_data['name'],
            'risk':          vendor_data['risk'],
            'request_count': vendor_data['request_count'],
            'post_count':    vendor_data['post_count'],
            'pii':           pii,
            'pii_count':     pii_count,
            'all_requests':  vendor_data['all_requests'],
        }

        if vendor_key == 'leadid':
            leadid_in_vendors = True
            leadid_request_count = vendor_data['request_count']

    # 7. Build LeadID timeline from the lightweight captures collected in step 3
    leadid_timeline = build_simple_timeline(leadid_captures) if leadid_captures else []

    # 8. Count total unique PII across all vendors
    all_pii = set()
    for vendor_data in vendors_with_pii.values():
        for pii in vendor_data['pii']:
            all_pii.add(f"{pii['type']}:{pii['value']}")

    return {
        'first_party':    first_party,
        'total_requests': total_requests,
        'post_requests':  post_requests,

        'vendors_detected': len(detected_vendors),
        'vendors':          vendors_with_pii,

        'first_party_pii': first_party_analysis,

        # Legacy LeadID fields
        'leadid_detected':       leadid_in_vendors,
        'leadid_request_count':  leadid_request_count,
        'timeline':              leadid_timeline,

        'pii_count':     len(all_pii),
        'session_start': session_start,
        'session_date':  session_date,
    }


def build_litigation_extract(results, risk):
    """
    Build a compact, structured text representation of the HAR analysis
    for use as the Claude API user message. Extracts everything Claude
    needs to produce a litigation package — no re-parsing of the HAR.

    Returns a plain-text string (~8-15K tokens) covering:
      - Domain, session date, consent detection
      - Full vendor timeline with timestamps, HAR indices, PII values
      - First-party PII collected
      - Identifier propagation across vendors
      - Damages (conservative / hybrid / aggressive)
    """
    vendors      = results.get('vendors', {})
    domain       = results.get('first_party', 'Unknown')
    session_date = results.get('session_date', 'Unknown')
    consent      = 'NO (no consent mechanism detected in session)'

    lines = []
    lines.append(f"Domain: {domain}")
    lines.append(f"Session Date: {session_date}")
    lines.append(f"Consent Detected: {consent}")
    lines.append(f"Risk Score: {risk['score']}/100 [{risk['grade_label']}]")
    lines.append(f"Estimated Damages: ${risk['estimated_damages']:,}")
    lines.append("")

    # ── Vendor timeline ───────────────────────────────────────────────────────
    lines.append("=" * 60)
    lines.append("VENDOR TIMELINE (chronological)")
    lines.append("=" * 60)

    # Collect all requests across all vendors, sorted by timestamp
    all_events = []
    for vkey, vdata in vendors.items():
        vname = vdata['name']
        for req in vdata.get('all_requests', []):
            all_events.append({
                'timestamp': req.get('timestamp', ''),
                'vendor':    vname,
                'method':    req.get('method', ''),
                'url':       req.get('url', ''),
                'has_pii':   req.get('has_pii', False),
            })
        for pii in vdata.get('pii', []):
            # PII items have their own timestamp and URL
            all_events.append({
                'timestamp': pii.get('timestamp', ''),
                'vendor':    vname,
                'method':    pii.get('request_details', {}).get('method', ''),
                'url':       pii.get('url', ''),
                'has_pii':   True,
                'pii_item':  pii,
            })

    # Sort by timestamp
    all_events.sort(key=lambda x: x.get('timestamp', ''))

    # Deduplicate — use (vendor, url, timestamp) as key, prefer pii_item entries
    seen_events = {}
    for ev in all_events:
        key = (ev['vendor'], ev['url'][:60], ev['timestamp'][:19])
        if key not in seen_events or ev.get('pii_item'):
            seen_events[key] = ev

    for idx, ev in enumerate(seen_events.values(), start=1):
        ts      = ev.get('timestamp', '')[:23]
        vendor  = ev['vendor']
        method  = ev.get('method', 'GET')
        url     = ev.get('url', '')
        pii     = ev.get('pii_item')

        lines.append("")
        lines.append(f"[{ts}] #{idx} {vendor} | {method} {url[:80]}")

        if pii:
            ptype = pii.get('type', '')
            is_plain  = 'Hashed' not in ptype
            is_hashed = 'Hashed' in ptype
            lines.append(f"  Type: {'PLAINTEXT PII' if is_plain else 'HASHED PII'}")
            lines.append(f"  Field: {pii.get('field', '')} → {pii.get('value', '')}")
            legal = pii.get('legal_context', {})
            if legal.get('cipa_elements'):
                lines.append(f"  Legal: {legal['cipa_elements'][0] if legal['cipa_elements'] else ''}")
        else:
            lines.append(f"  Type: Session/Identifier/Tracking")

    lines.append("")

    # ── First-party PII ───────────────────────────────────────────────────────
    fp = results.get('first_party_pii', {})
    fp_items = fp.get('pii_items', [])
    if fp_items:
        lines.append("=" * 60)
        lines.append("FIRST-PARTY PII COLLECTED (form fields submitted)")
        lines.append("=" * 60)
        for item in fp_items:
            lines.append(f"  • {item.get('type', '')}: {item.get('value', '')}")
        lines.append("")

    # ── Vendor PII summary ────────────────────────────────────────────────────
    # Pre-compute verified hash mappings.
    # For each first-party plaintext value, compute SHA-256 / SHA-1 / MD5
    # and check whether any third-party hash matches. Only flag as verified
    # if the cryptographic check passes — never assert linkage otherwise.
    import hashlib

    def _hashes_of(plaintext):
        """Return set of common hash representations for a plaintext value."""
        encoded = plaintext.encode('utf-8')
        return {
            hashlib.sha256(encoded).hexdigest(),
            hashlib.sha1(encoded).hexdigest(),
            hashlib.md5(encoded).hexdigest(),
            # Phone numbers are sometimes normalised before hashing
            hashlib.sha256(plaintext.replace(' ', '').replace('-', '').encode()).hexdigest(),
            hashlib.sha256(('+1' + plaintext).encode()).hexdigest(),
        }

    # Build a map: hash_value → (plaintext_value, plaintext_type)
    verified_hash_map = {}
    for item in fp_items:
        for h in _hashes_of(item.get('value', '')):
            verified_hash_map[h] = (item.get('value', ''), item.get('type', ''))

    lines.append("=" * 60)
    lines.append("THIRD-PARTY PII TRANSMISSIONS")
    lines.append("=" * 60)
    for vkey, vdata in vendors.items():
        pii_list = vdata.get('pii', [])
        if not pii_list:
            continue
        lines.append(f"\nVendor: {vdata['name']}")
        lines.append(f"  Requests: {vdata['request_count']}  |  PII items: {vdata['pii_count']}")
        for pii in pii_list:
            ptype  = pii.get('type', '')
            pval   = pii.get('value', '')
            field  = pii.get('field', '')
            ts     = pii.get('timestamp', '')[:23]
            method = pii.get('request_details', {}).get('method', '')

            # Check if this is a hash and whether we can verify it
            is_hash = ptype.startswith('Hashed')
            verified_plaintext = None
            verified_type      = None
            if is_hash:
                # pval may be truncated (e.g. "db8d59b2...") — look for full value
                # The full value is in the field name or other pii entries
                full_hash = pval.rstrip('.')
                if full_hash in verified_hash_map:
                    verified_plaintext, verified_type = verified_hash_map[full_hash]

            line = f"  [{ts}] {method} | {field} = {pval}  [{ptype}]"
            if is_hash and verified_plaintext:
                line += f"  ✓ VERIFIED: matches first-party {verified_type} value"
            elif is_hash:
                line += f"  — UNVERIFIED: hash does not match any known first-party plaintext value"
            lines.append(line)
    lines.append("")

    # ── Identifier propagation ────────────────────────────────────────────────
    # Find values that appear across multiple vendors
    value_to_vendors = {}
    for vdata in vendors.values():
        vname = vdata['name']
        for pii in vdata.get('pii', []):
            val = pii.get('value', '')
            if val and len(val) > 4:
                if val not in value_to_vendors:
                    value_to_vendors[val] = set()
                value_to_vendors[val].add(vname)

    cross_vendor = {v: vs for v, vs in value_to_vendors.items() if len(vs) > 1}
    if cross_vendor:
        lines.append("=" * 60)
        lines.append("IDENTIFIER PROPAGATION (values seen across multiple vendors)")
        lines.append("=" * 60)
        for val, vendor_set in cross_vendor.items():
            lines.append(f"  Value: {val[:64]}")
            lines.append(f"  Seen by: {', '.join(vendor_set)}")
        lines.append("")

    # ── Damages ───────────────────────────────────────────────────────────────
    lines.append("=" * 60)
    lines.append("DAMAGES ANALYSIS")
    lines.append("=" * 60)
    bk = risk.get('breakdown', {})
    lines.append(f"Score breakdown: Plaintext={bk.get('plaintext',0)} | "
                 f"Hashed={bk.get('hashed',0)} | Damages={bk.get('damages',0)} | "
                 f"Vendors={bk.get('vendors',0)} | Aggravating={bk.get('aggravating',0)}")

    # Damages tiers (using score-based estimates)
    score = risk['score']
    base  = risk['estimated_damages']
    conservative_val = base
    hybrid_val       = min(base * 4,  500000)
    aggressive_val   = min(base * 10, 1000000)

    lines.append(f"  Conservative: ${conservative_val:,}")
    lines.append(f"  Hybrid:       ${hybrid_val:,}")
    lines.append(f"  Aggressive:   ${aggressive_val:,}")
    lines.append("")
    lines.append("LeadID/Jornaya detected: " + ("YES" if results.get('leadid_detected') else "NO"))
    lines.append(f"Total requests in session: {results.get('total_requests', 0)}")

    return '\n'.join(lines)


def save_single_analysis(filename, results, risk, litigation_extract=None):
    """
    Save a completed single-file analysis summary to MongoDB.
    Stores the litigation_extract text so it can be retrieved for
    report generation without re-analyzing the HAR file.
    Non-fatal: if MongoDB is unavailable, logs and returns None.
    Returns the analysis_id string so the frontend can link to it.
    """
    if not MONGO_ENABLED:
        return None
    try:
        analysis_id = str(uuid.uuid4())
        vendors = results.get('vendors', {})

        vendor_pii_summary = [
            {
                'vendor':     v['name'],
                'pii_count':  v.get('pii_count', 0),
                'pii_types':  list({p['type'] for p in v.get('pii', [])}),
            }
            for v in vendors.values()
            if v.get('pii_count', 0) > 0
        ]

        doc = {
            'analysis_id':           analysis_id,
            'created_at':            datetime.now(timezone.utc),
            'filename':              filename,
            'domain':                results.get('first_party', 'Unknown'),
            'total_requests':        results.get('total_requests', 0),
            'vendors_detected':      results.get('vendors_detected', 0),
            'leadid_detected':       results.get('leadid_detected', False),
            'repair_used':           results.get('repair_used', False),
            'first_party_pii_count': results.get('first_party_pii', {}).get('pii_count', 0),
            'vendor_pii_count':      sum(v.get('pii_count', 0) for v in vendors.values()),
            'vendor_pii_summary':    vendor_pii_summary,
            # Risk score fields
            'score':                 risk['score'],
            'grade':                 risk['grade'],
            'grade_label':           risk['grade_label'],
            'estimated_damages':     risk['estimated_damages'],
            'top_violations':        risk['top_violations'],
            'breakdown':             risk['breakdown'],
            # Phase 3: litigation package fields
            'litigation_extract':    litigation_extract,
            'has_report':            False,
            'report_id':             None,
        }

        _col_single.insert_one(doc)
        print(f"✅ Single analysis saved to MongoDB: {analysis_id}")
        return analysis_id

    except Exception as e:
        print(f"⚠️  MongoDB save failed (single): {e}")
        return None


def save_bulk_ranking(summaries):
    """
    Save a completed bulk ranking session to MongoDB.
    Non-fatal: if MongoDB is unavailable, logs and returns None.
    Returns the session_id string so the frontend can link to it.
    """
    if not MONGO_ENABLED:
        return None
    try:
        session_id = str(uuid.uuid4())

        # Only store completed files, sorted by score descending
        completed = [s for s in summaries if s.get('status') == 'complete']
        completed.sort(key=lambda x: x.get('score', 0), reverse=True)

        ranked_results = [
            {
                'rank':              i + 1,
                'filename':          s.get('filename', ''),
                'domain':            s.get('domain', 'Unknown'),
                'score':             s.get('score', 0),
                'grade':             s.get('grade', 'lower'),
                'grade_label':       s.get('grade_label', 'Lower Risk'),
                'estimated_damages': s.get('estimated_damages', 0),
                'top_violations':    s.get('top_violations', []),
                'breakdown':         s.get('breakdown', {}),
                'leadid_detected':   s.get('leadid_detected', False),
                'total_requests':    s.get('total_entries', 0),
                'vendors_detected':  s.get('vendors_detected', 0),
                'vendor_pii_count':  s.get('vendor_pii_count', 0),
            }
            for i, s in enumerate(completed)
        ]

        doc = {
            'session_id':     session_id,
            'created_at':     datetime.now(timezone.utc),
            'files_analyzed': len(completed),
            'files_errored':  len(summaries) - len(completed),
            'ranked_results': ranked_results,
        }

        _col_bulk.insert_one(doc)
        print(f"✅ Bulk ranking saved to MongoDB: {session_id}")
        return session_id

    except Exception as e:
        print(f"⚠️  MongoDB save failed (bulk): {e}")
        return None


@app.route('/')
def index():
    """Serve the main HTML page"""
    return send_from_directory('static', 'index.html')

@app.route('/bulk')
def bulk():
    """Serve the bulk ranker page"""
    return send_from_directory('static', 'bulk.html')

@app.route('/history')
def history():
    """Serve the analysis history page"""
    return send_from_directory('static', 'history.html')

@app.route('/admin')
def admin():
    """Serve the admin panel page"""
    return send_from_directory('static', 'admin.html')

# ── Admin settings helpers ────────────────────────────────────────────────────

DEFAULT_SYSTEM_PROMPT = """You are a privacy law analyst operating in forensic reconstruction mode.

You will receive structured HAR analysis data extracted from a real browser session. Your task is to produce a litigation-ready package following the exact format and section structure provided.

CORE RULES:
- Use only the provided HAR analysis data as your source of truth
- Do not invent, infer, or carry over data from prior sessions
- Every factual claim must cite a specific HAR index, field name, or value from the data
- Do not assume consent unless the data explicitly shows a consent mechanism
- Operate at expert witness level — courtroom-ready language throughout

OUTPUT FORMAT:
Produce all 7 sections in order:
1. Declaration-Ready Timeline (pre/post consent boundary, tabular format)
2. Detailed Expert Declaration (numbered paragraphs, formal language)
3. CIPA Elements Table (§§ 631 and 632, tabular)
4. Vendor-by-Vendor Data Classification Matrix (tabular)
5. Identifier Propagation Map (tabular)
6. Partial-to-Full Value Reconstruction (tabular)
7. Damages Analysis (conservative / hybrid / aggressive, tabular with HAR IDs)

Format output as clean Markdown with proper table syntax.

CRITICAL RULE — HASH VERIFICATION:
The analysis data marks each hashed PII value as either:
  - "VERIFIED: matches first-party [type] value" — cryptographic match confirmed
  - "UNVERIFIED: hash does not match any known first-party plaintext value"

You MUST follow these rules strictly:
- Only assert that a hash corresponds to a specific plaintext value if it is marked VERIFIED
- If a hash is marked UNVERIFIED, state only that a hashed value was transmitted — do not claim to know what plaintext it represents
- Never infer or assume a hash-to-plaintext linkage based on proximity of data points alone
- In the damages analysis, unverified hashes may still be counted as violations but must be described accurately as "unverified hashed transmission" not as a confirmed disclosure of a specific value"""

DEFAULT_MODEL = 'claude-haiku-4-5-20251001'

# Available models — updated via /api/admin/models/refresh
DEFAULT_MODEL_LIST = [
    {
        'id':          'claude-haiku-4-5-20251001',
        'name':        'Claude Haiku 4.5',
        'description': 'Fastest · Best for testing and drafts',
        'input_cost':  1.00,
        'output_cost': 5.00,
    },
    {
        'id':          'claude-sonnet-4-6',
        'name':        'Claude Sonnet 4.6',
        'description': 'Balanced · Recommended for production',
        'input_cost':  3.00,
        'output_cost': 15.00,
    },
    {
        'id':          'claude-opus-4-6',
        'name':        'Claude Opus 4.6',
        'description': 'Most capable · Court-ready final packages',
        'input_cost':  5.00,
        'output_cost': 25.00,
    },
]


def get_admin_settings():
    """
    Return current admin settings from MongoDB.
    Falls back to defaults if DB unavailable or doc not found.
    """
    defaults = {
        'system_prompt':      DEFAULT_SYSTEM_PROMPT,
        'prompt_version':     'v1.0',
        'blank_template':     '',
        'sample_reference':   '',
        'default_model':      DEFAULT_MODEL,
        'model_list':         DEFAULT_MODEL_LIST,
        'models_refreshed_at': None,
    }
    if not MONGO_ENABLED:
        return defaults
    try:
        doc = _col_settings.find_one({'_id': 'global'})
        if not doc:
            return defaults
        doc.pop('_id', None)
        # Fill any missing keys with defaults
        for k, v in defaults.items():
            if k not in doc:
                doc[k] = v
        return doc
    except Exception as e:
        print(f'⚠️  get_admin_settings failed: {e}')
        return defaults


# ── Admin API routes ──────────────────────────────────────────────────────────

@app.route('/api/admin/settings', methods=['GET'])
def api_admin_settings_get():
    """Return current admin settings (prompt, template, default model, model list)."""
    settings = get_admin_settings()
    # Convert datetime to string if present
    if settings.get('models_refreshed_at') and not isinstance(settings['models_refreshed_at'], str):
        settings['models_refreshed_at'] = settings['models_refreshed_at'].isoformat()
    return jsonify(settings)


@app.route('/api/admin/settings', methods=['POST'])
def api_admin_settings_post():
    """
    Save admin settings to MongoDB.
    Accepts partial updates — only fields present in the request body are updated.
    """
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available', 'saved': False}), 503

    data = request.get_json(silent=True) or {}
    allowed = {'system_prompt', 'blank_template', 'sample_reference', 'default_model'}
    update  = {k: v for k, v in data.items() if k in allowed}

    if not update:
        return jsonify({'error': 'No valid fields provided', 'saved': False}), 400

    # Auto-increment prompt version when system_prompt changes
    if 'system_prompt' in update:
        try:
            current = _col_settings.find_one({'_id': 'global'}) or {}
            ver_str = current.get('prompt_version', 'v1.0')
            major, minor = ver_str.lstrip('v').split('.')
            update['prompt_version'] = f"v{major}.{int(minor) + 1}"
        except Exception:
            update['prompt_version'] = 'v1.1'

    update['last_updated'] = datetime.now(timezone.utc)

    try:
        _col_settings.update_one(
            {'_id': 'global'},
            {'$set': update},
            upsert=True
        )
        return jsonify({'saved': True, 'updated_fields': list(update.keys())})
    except Exception as e:
        print(f'⚠️  admin settings save failed: {e}')
        return jsonify({'error': str(e), 'saved': False}), 500


@app.route('/api/admin/settings/upload', methods=['POST'])
def api_admin_settings_upload():
    """
    Accept a file upload for blank_template or sample_reference.
    Supported formats: .txt, .md, .pdf (text extracted via pypdf).
    Content is saved as text to the admin_settings MongoDB document.
    """
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available', 'saved': False}), 503

    field = request.form.get('field')  # 'blank_template' or 'sample_reference'
    if field not in ('blank_template', 'sample_reference'):
        return jsonify({'error': 'Invalid field name', 'saved': False}), 400

    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'error': 'No file provided', 'saved': False}), 400

    fname = file.filename.lower()
    allowed = ('.txt', '.md', '.pdf')
    if not any(fname.endswith(ext) for ext in allowed):
        return jsonify({
            'error': f'Only {", ".join(allowed)} files are supported',
            'saved': False
        }), 400

    try:
        raw_bytes = file.read()
        if fname.endswith('.pdf'):
            # Extract text from PDF using pypdf (already installed)
            import io
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                pages  = [page.extract_text() or '' for page in reader.pages]
                text   = '\n\n'.join(pages).strip()
                if not text:
                    return jsonify({
                        'error': 'PDF appears to be scanned or image-only — no extractable text found. Please use a text-based PDF.',
                        'saved': False
                    }), 400
            except Exception as pdf_err:
                return jsonify({'error': f'Could not read PDF: {pdf_err}', 'saved': False}), 400
        else:
            text = raw_bytes.decode('utf-8', errors='replace')
    except Exception as e:
        return jsonify({'error': f'Could not read file: {e}', 'saved': False}), 400

    try:
        _col_settings.update_one(
            {'_id': 'global'},
            {'$set': {
                field:          text,
                'last_updated': datetime.now(timezone.utc),
            }},
            upsert=True
        )
        return jsonify({
            'saved':    True,
            'field':    field,
            'filename': file.filename,
            'length':   len(text),
        })
    except Exception as e:
        print(f'⚠️  file upload save failed: {e}')
        return jsonify({'error': str(e), 'saved': False}), 500


@app.route('/api/admin/models/refresh', methods=['POST'])
def api_admin_models_refresh():
    """
    Refresh the stored model list.
    Phase 2b: returns the hardcoded default list (no live Anthropic API call).
    Phase 3 update: replace with a live call to the Anthropic models endpoint.
    """
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available', 'refreshed': False}), 503

    try:
        now = datetime.now(timezone.utc)
        _col_settings.update_one(
            {'_id': 'global'},
            {'$set': {
                'model_list':          DEFAULT_MODEL_LIST,
                'models_refreshed_at': now,
                'last_updated':        now,
            }},
            upsert=True
        )
        return jsonify({
            'refreshed':    True,
            'model_count':  len(DEFAULT_MODEL_LIST),
            'refreshed_at': now.isoformat(),
            'models':       DEFAULT_MODEL_LIST,
        })
    except Exception as e:
        print(f'⚠️  model refresh failed: {e}')
        return jsonify({'error': str(e), 'refreshed': False}), 500


@app.route('/api/admin/test-generate', methods=['POST'])
def api_admin_test_generate():
    """
    Test generation endpoint — Phase 2b stub.
    Returns a structured HTML report using synthetic data and the saved prompt.
    No Claude API call is made. This validates the full pipeline flow
    (settings load → prompt format → HTML render → download) without spending tokens.
    Phase 3 will replace the stub response with a real Claude API call.
    """
    data     = request.get_json(silent=True) or {}
    model    = data.get('model', DEFAULT_MODEL)
    settings = get_admin_settings()

    # Synthetic test extract — mirrors what build_litigation_extract() will produce
    test_extract = {
        'domain':             'www.example-insurance.com',
        'session_date':       '2026-04-09',
        'consent_detected':   False,
        'vendor_timeline': [
            {
                'index': 38, 'timestamp': '2026-04-09T21:05:22.076Z',
                'vendor': 'LeadID/Jornaya', 'method': 'GET',
                'url': 'https://i.leadid.com/init',
                'type': 'session_init',
                'identifiers': {'pid': 'abc123', 'token': 'tok_xyz', 'lck': 'lck_def'},
                'pii': [], 'pre_consent': True,
                'legal_note': 'Session initialization before any consent boundary',
            },
            {
                'index': 64, 'timestamp': '2026-04-09T21:05:33.453Z',
                'vendor': 'LeadID/Jornaya', 'method': 'POST',
                'url': 'https://track.leadid.com/track',
                'type': 'plaintext_pii',
                'identifiers': {},
                'pii': [{'field': 'zipcode', 'value': '95050', 'type': 'Zip Code'}],
                'pre_consent': True,
                'legal_note': 'Plaintext PII captured before form submission §631(a)',
            },
            {
                'index': 124, 'timestamp': '2026-04-09T21:06:08.818Z',
                'vendor': 'LeadID/Jornaya', 'method': 'POST',
                'url': 'https://track.leadid.com/track',
                'type': 'plaintext_pii',
                'identifiers': {},
                'pii': [{'field': 'contactName', 'value': 'Matthew Vargas', 'type': 'Full Name'}],
                'pre_consent': False,
                'legal_note': 'Plaintext full name captured post-transition §632',
            },
            {
                'index': 131, 'timestamp': '2026-04-09T21:06:25.881Z',
                'vendor': 'LeadID/Jornaya', 'method': 'POST',
                'url': 'https://track.leadid.com/track',
                'type': 'plaintext_pii',
                'identifiers': {},
                'pii': [{'field': 'email', 'value': 'matthew@example.com', 'type': 'Email'}],
                'pre_consent': False,
                'legal_note': 'Plaintext email captured post-transition §632',
            },
            {
                'index': 132, 'timestamp': '2026-04-09T21:06:50.018Z',
                'vendor': 'LeadID/Jornaya', 'method': 'POST',
                'url': 'https://track.leadid.com/track',
                'type': 'plaintext_pii',
                'identifiers': {},
                'pii': [{'field': 'contactPhone', 'value': '6123277188', 'type': 'Phone'}],
                'pre_consent': False,
                'legal_note': 'Plaintext phone captured post-transition §632',
            },
            {
                'index': 287, 'timestamp': '2026-04-09T21:35:04.391Z',
                'vendor': 'TikTok Pixel', 'method': 'POST',
                'url': 'https://analytics.tiktok.com/api/v2/pixel/track/',
                'type': 'hashed_pii',
                'identifiers': {},
                'pii': [{'field': 'context.user.eb_phone_number',
                          'value': 'db8d59b22475fb52ce94321a25470711ca7920b6d1e86e65b5f052b0eb04ed52',
                          'type': 'Hashed Phone (SHA-256)'}],
                'pre_consent': False,
                'legal_note': 'SHA-256 hashed phone enables cross-site user identification',
            },
        ],
        'damages_conservative': {'count': 7,  'value': 35000},
        'damages_hybrid':       {'count': 29, 'value': 145000},
        'damages_aggressive':   {'count': 79, 'value': 395000},
    }

    # Build stub markdown response — same structure Claude will return in Phase 3
    stub_md = build_stub_report(test_extract, model, settings.get('prompt_version', 'v1.0'))

    # Convert to HTML
    html_content = markdown_to_html(stub_md)

    return jsonify({
        'status':         'complete',
        'stub':           True,
        'model':          model,
        'prompt_version': settings.get('prompt_version', 'v1.0'),
        'html':           html_content,
        'input_tokens':   0,
        'output_tokens':  0,
    })


def build_stub_report(extract, model, prompt_version):
    """
    Build a stub litigation package in Markdown from the test extract.
    This is the same structure Claude will produce in Phase 3.
    """
    domain   = extract['domain']
    date_str = extract['session_date']
    timeline = extract['vendor_timeline']
    consent  = 'YES' if extract['consent_detected'] else 'NO'

    # Split timeline pre/post consent
    pre  = [e for e in timeline if e['pre_consent']]
    post = [e for e in timeline if not e['pre_consent']]

    def timeline_rows(events):
        rows = []
        for e in events:
            pii_str = ', '.join(f"{p['field']}={p['value']}" for p in e['pii']) if e['pii'] else '—'
            id_str  = ', '.join(f"{k}={v}" for k, v in e.get('identifiers', {}).items()) or '—'
            data    = pii_str if pii_str != '—' else id_str
            rows.append(f"| {e['timestamp']} | #{e['index']} | {e['vendor']} | "
                        f"{'PLAINTEXT PII' if e['type']=='plaintext_pii' else 'Hashed PII' if e['type']=='hashed_pii' else 'Session/Identifiers'} | "
                        f"{data} | {e['legal_note']} |")
        return '\n'.join(rows) if rows else '| — | — | — | — | — | — |'

    # Vendor matrix
    vendors = {}
    for e in timeline:
        v = e['vendor']
        if v not in vendors:
            vendors[v] = {'timing': [], 'types': set(), 'data': [], 'ids': []}
        vendors[v]['timing'].append(f"#{e['index']}")
        vendors[v]['types'].add(e['type'])
        for p in e.get('pii', []):
            vendors[v]['data'].append(f"{p['field']}={p['value']}")
        vendors[v]['ids'].append(str(e['index']))

    vendor_rows = []
    for vname, vdata in vendors.items():
        timing   = ', '.join(vdata['timing'])
        types    = ', '.join(vdata['types'])
        data_str = '; '.join(vdata['data']) if vdata['data'] else 'Identifiers only'
        ids_str  = ', '.join(vdata['ids'])
        vendor_rows.append(f"| {vname} | {timing} | {types} | {data_str} | {ids_str} |")

    dmg = extract
    con = dmg['damages_conservative']
    hyb = dmg['damages_hybrid']
    agg = dmg['damages_aggressive']

    return f"""# Final Litigation Package — {domain}
## ⚠️ STUB OUTPUT — Phase 2b test (no Claude API call made)
**Domain:** {domain}  |  **Session Date:** {date_str}  |  **Consent Detected:** {consent}
**Model:** {model}  |  **Prompt Version:** {prompt_version}

---

## 1. Declaration-Ready Timeline

### A. Pre-Consent Events

| Time (UTC) | HAR ID | Vendor | Data Type | Transmitted | Legal Significance |
|---|---|---|---|---|---|
{timeline_rows(pre)}

### B. Post-Consent Events

| Time (UTC) | HAR ID | Vendor | Data Type | Transmitted | Legal Significance |
|---|---|---|---|---|---|
{timeline_rows(post)}

---

## 2. Expert Declaration

I, [Name], declare as follows:

1. I reviewed structured HAR analysis data for {domain} captured on {date_str}.
2. I did not identify a consent management platform or affirmative consent event in the session data.
3. LeadID/Jornaya received plaintext user input — including zip code, full name, email address, and phone number — both before and after the first meaningful form step transition.
4. TikTok Pixel received a SHA-256 hashed phone number at HAR #287 enabling cross-site user identification.
5. No verified consent mechanism was observed preceding any of the third-party data transmissions described above.

I declare under penalty of perjury under the laws of the State of California that the foregoing is true and correct.

---

## 3. CIPA Elements (§§ 631 and 632)

| Element | Requirement | Best Evidence | Why It Proves the Element |
|---|---|---|---|
| Interception during transmission §631(a) | Third party reads communication in transit | HAR #64: zipcode=95050 captured pre-submission | LeadID received plaintext data before the user submitted the form |
| Contents of a communication §631(a) | Substance, not just routing information | HAR #124, #131, #132: name, email, phone | Actual typed field values — not anonymous identifiers |
| Third-party disclosure §632 | Disclosed to third party without authorization | All LeadID SaveFormField requests post-transition | No consent mechanism observed; data disclosed outward |
| No valid consent §§631–632 | Interception occurs before or without consent | No CMP or acceptance event in session data | Session data lacks any conventional consent flow |

---

## 4. Vendor-by-Vendor Data Classification Matrix

| Vendor | HAR IDs | Data Type | Exactly What Received | Supporting IDs |
|---|---|---|---|---|
{chr(10).join(vendor_rows)}

---

## 5. Identifier Propagation Map

| Recurring Identifier | Where It Appears | Vendors | Why It Matters |
|---|---|---|---|
| pid=abc123 | HAR #38 init + all SaveFormField requests | LeadID/Jornaya | Ties full session together from initialization through field captures |
| SHA-256 phone hash | HAR #287 TikTok payload | TikTok Pixel | Enables cross-site tracking and persistent user re-identification |

---

## 6. Partial-to-Full Value Reconstruction

| Field | Partial Values Observed | Full Final Value | HAR IDs | Significance |
|---|---|---|---|---|
| Zip Code | No partial sequence | 95050 | #64 | Plaintext final value captured by LeadID pre-submission |
| Full Name | No partial sequence | Matthew Vargas | #124 | Plaintext full name captured by LeadID |
| Email | No partial sequence | matthew@example.com | #131 | Plaintext email captured by LeadID |
| Phone | No partial sequence | 6123277188 | #132 | Plaintext phone captured by LeadID |
| Hashed Phone | N/A — hash only | db8d59b2... (SHA-256) | #287 | Hashed value received by TikTok enabling re-identification |

---

## 7. Damages Analysis

| Approach | Counting Logic | HAR IDs Included | Count | Statutory Value |
|---|---|---|---|---|
| Conservative | Strongest plaintext PII captures only | #64, #124, #131, #132 | {con['count']} | ${con['value']:,} |
| Hybrid | Conservative + corroborating identifier transmissions | #38, #64, #124, #131, #132, #287 | {hyb['count']} | ${hyb['value']:,} |
| Aggressive | All relevant third-party transmissions | All vendor requests | {agg['count']} | ${agg['value']:,} |

**Recommended approach:** Conservative is the cleanest and most defensible. Hybrid adds surrounding context without over-counting repetitive calls.
"""


def markdown_to_html(md):
    """
    Convert Markdown litigation package to styled HTML.
    Handles: headings, tables (with thead/tbody), bold/italic,
    numbered lists, horizontal rules, and warning blocks.
    Used for both stub and real Claude output.
    """
    lines  = md.split('\n')
    output = []
    i      = 0

    while i < len(lines):
        line = lines[i]

        # Markdown table — collect all consecutive pipe lines
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            tbl  = '<table>\n'
            first = True
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                # Skip separator rows
                if all(re.match(r'^-+$', c.replace(':', '').strip()) for c in cells if c):
                    continue
                if first:
                    tbl  += '<thead><tr>' + ''.join(f'<th>{c}</th>' for c in cells) + '</tr></thead>\n<tbody>\n'
                    first = False
                else:
                    tbl  += '<tr>' + ''.join(f'<td>{inline(c)}</td>' for c in cells) + '</tr>\n'
            tbl += '</tbody></table>\n'
            output.append(tbl)
            continue

        # Warning block
        if line.startswith('## ⚠️'):
            output.append(f'<div class=\"warning\"><strong>⚠️</strong> {inline(line[6:])}</div>')
            i += 1; continue

        # Headings
        if line.startswith('### '): output.append(f'<h3>{inline(line[4:])}</h3>'); i += 1; continue
        if line.startswith('## '):  output.append(f'<h2>{inline(line[3:])}</h2>'); i += 1; continue
        if line.startswith('# '):   output.append(f'<h1>{inline(line[2:])}</h1>'); i += 1; continue

        # Horizontal rule
        if line.strip() == '---': output.append('<hr>'); i += 1; continue

        # Numbered list
        if re.match(r'^\d+\.\s', line):
            output.append(f'<li>{inline(line[line.index(".")+2:])}</li>')
            i += 1; continue

        # Blank line
        if line.strip() == '': i += 1; continue

        # Paragraph
        output.append(f'<p>{inline(line)}</p>')
        i += 1

    body = '\n'.join(output)

    # Wrap consecutive <li> in <ol>
    body = re.sub(r'((?:<li>.*?</li>\n?)+)', lambda m: '<ol>\n' + m.group(0) + '</ol>\n', body)

    return f"""<!DOCTYPE html>
<html lang=\"en\">
<head>
<meta charset=\"UTF-8\">
<title>Litigation Package</title>
<style>
  body {{ font-family: Arial, sans-serif; max-width: 980px; margin: 40px auto; padding: 0 32px; color: #222; line-height: 1.7; }}
  h1 {{ font-size: 1.4em; border-bottom: 2px solid #333; padding-bottom: 8px; margin-top: 32px; }}
  h2 {{ font-size: 1.15em; margin-top: 32px; color: #333; border-left: 3px solid #666; padding-left: 10px; }}
  h3 {{ font-size: 1em; margin-top: 20px; color: #444; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0 24px; font-size: 0.87em; }}
  thead th {{ background: #e8e8e8; padding: 9px 12px; border: 1px solid #bbb; text-align: left; font-weight: 600; }}
  tbody td {{ padding: 8px 12px; border: 1px solid #ddd; vertical-align: top; }}
  tbody tr:nth-child(even) {{ background: #f9f9f9; }}
  hr {{ border: none; border-top: 1px solid #e0e0e0; margin: 28px 0; }}
  p {{ margin: 8px 0; }}
  ol {{ margin: 8px 0 8px 24px; padding: 0; }}
  li {{ margin: 4px 0; }}
  code {{ background: #f4f4f4; padding: 1px 5px; border-radius: 3px; font-size: 0.9em; }}
  strong {{ color: #111; }}
  .warning {{ background: #fff8e1; border: 1px solid #f0c040; border-left: 4px solid #e6ac00;
              padding: 10px 16px; border-radius: 4px; font-size: 0.88em; margin: 12px 0; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def inline(text):
    """Apply inline markdown formatting (bold, italic, code)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*(.+?)\*',     r'<em>\1</em>',         text)
    text = re.sub(r'`(.+?)`',        r'<code>\1</code>',     text)
    return text


def inline_plain(text):
    """Strip markdown formatting to plain text (for docx runs)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',        r'\1', text)
    return text


def markdown_to_docx(md):
    """
    Convert Markdown litigation package to a styled .docx file.
    Returns a BytesIO object ready to serve as a file download.

    Handles:
      - H1/H2/H3 → Word heading styles
      - Tables    → Word tables with header row shading
      - Bold/italic inline formatting
      - Horizontal rules → paragraph spacing
      - Numbered lists → ListNumber style
      - Warning blocks (## ⚠️) → highlighted paragraph
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io, copy
    except ImportError:
        raise ImportError('python-docx is not installed — run: pip install python-docx==1.1.2')

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Heading styles ────────────────────────────────────────────────────────
    def set_heading_style(para, level):
        style_name = f'Heading {level}'
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            para.style = doc.styles['Normal']
        run = para.runs[0] if para.runs else para.add_run()
        if level == 1:
            run.font.size  = Pt(16)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        elif level == 2:
            run.font.size  = Pt(13)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif level == 3:
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Add paragraph with inline formatting ──────────────────────────────────
    def add_formatted_para(text, style='Normal', bold_all=False):
        para  = doc.add_paragraph(style=style)
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            else:
                run = para.add_run(part)
            if bold_all:
                run.bold = True
        return para

    # ── Add table ─────────────────────────────────────────────────────────────
    def add_table(header_cells, rows):
        col_count = len(header_cells)
        if col_count == 0:
            return
        tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
        tbl.style = 'Table Grid'

        # Header row
        hdr_row = tbl.rows[0]
        for i, cell_text in enumerate(header_cells):
            cell = hdr_row.cells[i]
            cell.text = inline_plain(cell_text)
            # Bold header text
            for run in cell.paragraphs[0].runs:
                run.bold = True
            # Light grey shading
            tc_pr = cell._tc.get_or_add_tcPr()
            shd   = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'E8E8E8')
            tc_pr.append(shd)

        # Data rows
        for r_idx, row_cells in enumerate(rows):
            row = tbl.rows[r_idx + 1]
            for c_idx, cell_text in enumerate(row_cells):
                if c_idx < col_count:
                    row.cells[c_idx].text = inline_plain(cell_text)

        doc.add_paragraph()  # spacing after table

    # ── Parse markdown line by line ───────────────────────────────────────────
    lines   = md.split('\n')
    i       = 0
    in_list = False

    while i < len(lines):
        line = lines[i]

        # Markdown table
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                rows_raw = []
                header   = None
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    if all(re.match(r'^-+$', c.replace(':', '').strip()) for c in cells if c):
                        continue  # skip separator
                    if header is None:
                        header = cells
                    else:
                        rows_raw.append(cells)
                if header:
                    add_table(header, rows_raw)
            continue

        # Warning block
        if line.startswith('## ⚠️') or line.startswith('## ⚠'):
            p = add_formatted_para(line[3:].strip(), 'Normal')
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x00) if p.runs else None
            i += 1; continue

        # Headings
        if line.startswith('# '):
            p = doc.add_heading(inline_plain(line[2:]), level=1)
            i += 1; continue
        if line.startswith('## '):
            p = doc.add_heading(inline_plain(line[3:]), level=2)
            i += 1; continue
        if line.startswith('### '):
            p = doc.add_heading(inline_plain(line[4:]), level=3)
            i += 1; continue

        # Horizontal rule — add spacing
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1; continue

        # Numbered list
        if re.match(r'^\d+\.\s', line):
            add_formatted_para(line[line.index('.')+2:], 'List Number')
            i += 1; continue

        # Blank line
        if line.strip() == '':
            i += 1; continue

        # Regular paragraph
        add_formatted_para(line)
        i += 1

    # Return as BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── History API endpoints ────────────────────────────────────────────────────

@app.route('/api/history/single', methods=['GET'])
def api_history_single():
    """Return list of past single-file analyses from MongoDB, newest first."""
    if not MONGO_ENABLED:
        return jsonify({'results': [], 'total': 0, 'db_enabled': False})
    try:
        docs = list(_col_single.find(
            {},
            {
                '_id': 0,
                'analysis_id': 1, 'created_at': 1, 'filename': 1, 'domain': 1,
                'total_requests': 1, 'vendors_detected': 1, 'leadid_detected': 1,
                'first_party_pii_count': 1, 'vendor_pii_count': 1,
                'score': 1, 'grade': 1, 'grade_label': 1,
                'estimated_damages': 1, 'top_violations': 1,
                'has_report': 1, 'report_id': 1,
            }
        ).sort('created_at', DESCENDING).limit(200))

        # Convert datetime to ISO string for JSON serialization
        for doc in docs:
            if 'created_at' in doc:
                doc['created_at'] = doc['created_at'].isoformat()

        return jsonify({'results': docs, 'total': len(docs), 'db_enabled': True})
    except Exception as e:
        print(f"⚠️  MongoDB query failed (single list): {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/history/single/<analysis_id>', methods=['GET'])
def api_history_single_detail(analysis_id):
    """Return full stored result for one single-file analysis."""
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503
    try:
        doc = _col_single.find_one(
            {'analysis_id': analysis_id},
            {'_id': 0}
        )
        if not doc:
            return jsonify({'error': 'Analysis not found'}), 404
        if 'created_at' in doc:
            doc['created_at'] = doc['created_at'].isoformat()
        return jsonify(doc)
    except Exception as e:
        print(f"⚠️  MongoDB query failed (single detail): {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/history/bulk', methods=['GET'])
def api_history_bulk():
    """Return list of past bulk ranking sessions from MongoDB, newest first."""
    if not MONGO_ENABLED:
        return jsonify({'results': [], 'total': 0, 'db_enabled': False})
    try:
        docs = list(_col_bulk.find(
            {},
            {
                '_id': 0,
                'session_id': 1, 'created_at': 1,
                'files_analyzed': 1, 'files_errored': 1,
                'ranked_results': 1,
            }
        ).sort('created_at', DESCENDING).limit(100))

        for doc in docs:
            if 'created_at' in doc:
                doc['created_at'] = doc['created_at'].isoformat()

        return jsonify({'results': docs, 'total': len(docs), 'db_enabled': True})
    except Exception as e:
        print(f"⚠️  MongoDB query failed (bulk list): {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/history/bulk/<session_id>', methods=['GET'])
def api_history_bulk_detail(session_id):
    """Return full stored result for one bulk ranking session."""
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503
    try:
        doc = _col_bulk.find_one(
            {'session_id': session_id},
            {'_id': 0}
        )
        if not doc:
            return jsonify({'error': 'Session not found'}), 404
        if 'created_at' in doc:
            doc['created_at'] = doc['created_at'].isoformat()
        return jsonify(doc)
    except Exception as e:
        print(f"⚠️  MongoDB query failed (bulk detail): {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/history/single/<analysis_id>', methods=['DELETE'])
def api_delete_single(analysis_id):
    """Delete a single-file analysis result from MongoDB."""
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503
    try:
        result = _col_single.delete_one({'analysis_id': analysis_id})
        if result.deleted_count == 0:
            return jsonify({'error': 'Analysis not found'}), 404
        print(f"🗑  Single analysis deleted: {analysis_id}")
        return jsonify({'deleted': True, 'analysis_id': analysis_id})
    except Exception as e:
        print(f"⚠️  MongoDB delete failed (single): {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/history/bulk/<session_id>', methods=['DELETE'])
def api_delete_bulk(session_id):
    """Delete a bulk ranking session from MongoDB."""
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503
    try:
        result = _col_bulk.delete_one({'session_id': session_id})
        if result.deleted_count == 0:
            return jsonify({'error': 'Session not found'}), 404
        print(f"🗑  Bulk session deleted: {session_id}")
        return jsonify({'deleted': True, 'session_id': session_id})
    except Exception as e:
        print(f"⚠️  MongoDB delete failed (bulk): {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/analyze', methods=['POST'])
def analyze():
    """API endpoint for analyzing HAR files"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.endswith('.har'):
            return jsonify({'error': 'File must be a .har file'}), 400
        
        # Load HAR data with robust error handling
        try:
            # Read file content
            file_content = file.read()
            
            # Decode with error handling for non-UTF-8 bytes
            try:
                har_text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                print("Warning: HAR file contains non-UTF-8 bytes, using replacement")
                har_text = file_content.decode('utf-8', errors='replace')
            
            # Remove control characters that break JSON parsing
            # Option C: free the original decoded string immediately so both
            # full-file strings never coexist in memory at once.
            har_text_clean = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', har_text)
            del har_text  # Option C: release ~file_size MB here

            # Try standard JSON parsing first
            repair_used = False
            repair_stats = None
            
            try:
                har_data = json.loads(har_text_clean)
                
                # Validate HAR structure
                if 'log' not in har_data or 'entries' not in har_data['log']:
                    raise ValueError('Invalid HAR file structure - missing log.entries')

                # Option B: strip response body text from parsed dict.
                # O(n entries) dict walk — much faster than pre-parse string manipulation.
                strip_response_bodies(har_data)
                del har_text_clean  # string no longer needed
                
            except (json.JSONDecodeError, ValueError) as e:
                # Standard parsing failed - use resilient parser
                print(f"Standard parsing failed: {e}")
                
                try:
                    entries, repair_stats = resilient_parse_har(har_text_clean)
                    
                    if not entries:
                        raise Exception("Resilient parser could not extract any valid entries")
                    
                    # Build a valid HAR structure
                    har_data = {
                        'log': {
                            'version': '1.2',
                            'creator': {
                                'name': 'HAR Analyzer with Repair',
                                'version': '1.0'
                            },
                            'entries': entries
                        }
                    }
                    
                    repair_used = True
                    
                except Exception as repair_error:
                    # Even resilient parsing failed
                    error_msg = (
                        f'HAR file is severely corrupted and could not be repaired. '
                        f'Error: {str(repair_error)}. '
                        f'Try re-exporting the HAR file from your browser.'
                    )
                    return jsonify({'error': error_msg}), 400
        
        except Exception as e:
            return jsonify({'error': f'Error reading file: {str(e)}'}), 400
        
        # Analyze the HAR data (whether from standard or resilient parsing)
        results = analyze_har_simple(har_data)
        
        # Add repair information to results
        if repair_used:
            results['repair_used'] = True
            results['repair_stats'] = repair_stats
            
            # HYBRID APPROACH: Always run plaintext scan when file was repaired
            # This supplements structured parsing with plaintext extraction
            print("File was repaired - running supplemental plaintext scan...")
            
            try:
                first_party_domain = results.get('first_party', '')
                plaintext_pii = analyze_har_as_plaintext(har_text_clean, first_party_domain)
                
                if plaintext_pii['pii_count'] > 0:
                    # Merge plaintext PII with structured parsing results
                    existing_pii = results.get('first_party_pii', {})
                    existing_items = existing_pii.get('pii_items', [])
                    plaintext_items = plaintext_pii.get('pii_items', [])
                    
                    # Combine and deduplicate
                    all_pii = existing_items + plaintext_items
                    
                    # Deduplicate by type:value
                    unique_pii = {}
                    for pii in all_pii:
                        key = f"{pii['type']}:{pii['value']}"
                        if key not in unique_pii:
                            unique_pii[key] = pii
                    
                    combined_items = list(unique_pii.values())
                    
                    # Update results with combined data
                    results['first_party_pii'] = {
                        'pii_items': combined_items,
                        'pii_count': len(combined_items),
                        'post_count': existing_pii.get('post_count', 0),
                        'method': 'hybrid' if existing_items else 'plaintext_fallback',
                        'structured_count': len(existing_items),
                        'plaintext_count': len(plaintext_items),
                        'note': f'Combined: {len(existing_items)} from parsed entries + {len(plaintext_items)} from plaintext scan'
                    }
                    
                    results['plaintext_supplement_used'] = True
                    print(f"✅ Hybrid scan: {len(existing_items)} structured + {len(plaintext_items)} plaintext = {len(combined_items)} total PII")
                else:
                    results['plaintext_supplement_used'] = False
            except Exception as e:
                print(f"Plaintext supplement error: {e}")
                results['plaintext_supplement_used'] = False
        else:
            results['repair_used'] = False
            results['plaintext_supplement_used'] = False
        
        # Save to MongoDB and attach analysis_id so frontend can link to history
        risk = compute_risk_score(results)

        # Build litigation extract for Phase 3 report generation
        try:
            litigation_extract = build_litigation_extract(results, risk)
        except Exception as ex:
            print(f"⚠️  build_litigation_extract failed (non-fatal): {ex}")
            litigation_extract = None

        analysis_id = save_single_analysis(file.filename, results, risk, litigation_extract)
        if analysis_id:
            results['analysis_id'] = analysis_id

        # Surface whether Claude is available so frontend can show/hide generate button
        results['claude_enabled'] = CLAUDE_ENABLED

        return jsonify(results)
    
    except Exception as e:
        print(f"Error analyzing HAR: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Error analyzing file: {str(e)}'}), 500



def compute_risk_score(results):
    """
    Compute a 0-100 privacy risk score using weighted sub-scores.

    Each signal is scored 0-100 on its own severity scale, then
    combined as a weighted average. This ensures even a single
    plaintext PII violation produces a meaningfully high score,
    rather than scoring against a theoretical maximum that real
    files almost never reach.

    Weights:
      40%  Plaintext PII to third parties  (most severe)
      20%  Hashed PII transmissions
      15%  Damages magnitude
      15%  Vendor exposure breadth
      10%  Aggravating factors (LeadID + GET with PII)

    Sub-score scales:
      Plaintext: 0=0, 1=72, 2=85, 3+=100
      Hashed:    0=0, 1=55, 2=75, 3+=100
      Damages:   $0=0, $5K=40, $10K=60, $25K=80, $40K+=100
      Vendors:   0=0, 1=45, 2=75, 3+=100
      Aggrav:    0 factors=0, 1=55, 2=100 (LeadID counts, GET PII counts)

    Grade thresholds:
      Critical  = any plaintext PII present, OR score >= 75
      High      = any hashed PII present,    OR score >= 50
      Elevated  = score >= 25
      Lower     = score <  25
    """
    vendors = results.get('vendors', {})

    # --- Deduplicate by type+value to avoid counting repeated requests ---
    seen_plain, seen_hashed, seen_dmg = set(), set(), set()
    plaintext_pii_items, hashed_pii_items = [], []
    total_damages = 0

    for v in vendors.values():
        for p in v.get('pii', []):
            key = f"{p.get('type','')}:{p.get('value','')}"
            is_hashed = p.get('type', '').startswith('Hashed')
            if is_hashed:
                if key not in seen_hashed:
                    seen_hashed.add(key)
                    hashed_pii_items.append(p)
            else:
                if key not in seen_plain:
                    seen_plain.add(key)
                    plaintext_pii_items.append(p)
            if key not in seen_dmg:
                seen_dmg.add(key)
                total_damages += p.get('legal_context', {}).get('estimated_damages', 0)

    vendors_with_pii = [v for v in vendors.values() if v.get('pii_count', 0) > 0]
    leadid_present   = results.get('leadid_detected', False)
    get_with_pii     = any(
        p.get('request_details', {}).get('method') == 'GET'
        for v in vendors.values()
        for p in v.get('pii', [])
    )

    # --- Sub-scores (0-100 each) ---
    n_plain  = len(plaintext_pii_items)
    n_hashed = len(hashed_pii_items)
    n_vendor = len(vendors_with_pii)
    n_aggrav = (1 if leadid_present else 0) + (1 if get_with_pii else 0)

    def _plain_sub(n):
        if n == 0: return 0
        if n == 1: return 72
        if n == 2: return 85
        return min(100, 85 + (n - 2) * 5)

    def _hashed_sub(n):
        if n == 0: return 0
        if n == 1: return 55
        if n == 2: return 75
        return min(100, 75 + (n - 2) * 10)

    def _damage_sub(d):
        if d == 0:      return 0
        if d <= 5000:   return 40
        if d <= 10000:  return 60
        if d <= 25000:  return 80
        return min(100, int((d / 40000) * 100))

    def _vendor_sub(n):
        if n == 0: return 0
        if n == 1: return 45
        if n == 2: return 75
        return min(100, 75 + (n - 2) * 10)

    def _aggrav_sub(n):
        return [0, 55, 100][min(n, 2)]

    sp = _plain_sub(n_plain)
    sh = _hashed_sub(n_hashed)
    sd = _damage_sub(total_damages)
    sv = _vendor_sub(n_vendor)
    sa = _aggrav_sub(n_aggrav)

    total_score = round(sp * 0.40 + sh * 0.20 + sd * 0.15 + sv * 0.15 + sa * 0.10)

    # --- Grade ---
    if plaintext_pii_items or total_score >= 75:
        grade, grade_label = 'critical', 'Critical Risk'
    elif hashed_pii_items or total_score >= 50:
        grade, grade_label = 'high', 'High Risk'
    elif total_score >= 25:
        grade, grade_label = 'elevated', 'Elevated Risk'
    else:
        grade, grade_label = 'lower', 'Lower Risk'

    # --- Top violations ---
    top_violations = []
    if plaintext_pii_items:
        unique_types = list(dict.fromkeys(p['type'] for p in plaintext_pii_items))
        top_violations.append(
            f"{n_plain} plaintext PII item{'' if n_plain == 1 else 's'} "
            f"sent to third parties ({', '.join(unique_types[:2])})"
        )
    if hashed_pii_items:
        top_violations.append(
            f"{n_hashed} unique hashed PII transmission{'' if n_hashed == 1 else 's'}"
        )
    if leadid_present:
        top_violations.append("LeadID/Jornaya intercepts form fields in real time")
    if get_with_pii:
        top_violations.append("PII exposed in GET request URLs (logged by servers)")
    if vendors_with_pii and not top_violations:
        names = [v['name'] for v in vendors_with_pii]
        top_violations.append(f"Vendors receiving data: {', '.join(names)}")

    return {
        'score':             total_score,
        'grade':             grade,
        'grade_label':       grade_label,
        'estimated_damages': total_damages,
        'top_violations':    top_violations[:3],
        'breakdown': {
            'plaintext': sp,
            'hashed':    sh,
            'damages':   sd,
            'vendors':   sv,
            'aggravating': sa,
        }
    }

@app.route('/analyze-bulk', methods=['POST'])
def analyze_bulk():
    """
    Sequential bulk analysis endpoint using Server-Sent Events (SSE).

    Memory strategy — save uploads to temp files immediately, then process
    one at a time from disk. This avoids holding all file contents in RAM
    simultaneously, which was the cause of OOM crashes on Render free tier.

    Previous approach: buffer all files as bytes in memory before streaming.
    Problem: decoding a 99MB file to a Python str (which uses 2x memory)
    while 4 other files' bytes are still in RAM spiked to 641MB+.

    Current approach:
      1. Stream each upload to a temp file on disk (64KB chunks, ~0 RAM)
      2. SSE generator opens each temp file one at a time, processes it,
         then deletes the temp file before moving to the next.
      3. Peak RAM = baseline + single file processing (~180MB for 5x100MB)
    """
    import tempfile, shutil

    MAX_FILES     = 5
    MAX_SIZE_BYTES = 100 * 1024 * 1024  # 100MB per file

    files = request.files.getlist('files')

    if not files or len(files) == 0:
        return jsonify({'error': 'No files uploaded'}), 400

    if len(files) > MAX_FILES:
        return jsonify({'error': f'Maximum {MAX_FILES} files allowed'}), 400

    # Phase 1: Stream each upload to a temp file — no full-file buffer in RAM
    temp_files = []   # list of (original_filename, temp_path, file_size)
    for f in files:
        fname = f.filename or 'unknown.har'
        tmp   = tempfile.NamedTemporaryFile(delete=False, suffix='.har')
        size  = 0
        try:
            for chunk in iter(lambda: f.stream.read(65536), b''):
                tmp.write(chunk)
                size += len(chunk)
        finally:
            tmp.close()
        temp_files.append((fname, tmp.name, size))

    def generate(temp_files):
        total     = len(temp_files)
        summaries = []

        for idx, (filename, tmp_path, file_size) in enumerate(temp_files):

            # Emit processing event immediately
            yield f"data: {json.dumps({'event': 'processing', 'index': idx, 'filename': filename, 'current': idx + 1, 'total': total})}\n\n"

            # --- Validation ---
            if not filename.endswith('.har'):
                result = {'event': 'file_done', 'index': idx, 'filename': filename,
                          'status': 'error', 'error': 'Not a .har file',
                          'current': idx + 1, 'total': total}
                summaries.append(result)
                os.unlink(tmp_path)
                yield f"data: {json.dumps(result)}\n\n"
                continue

            if file_size > MAX_SIZE_BYTES:
                size_mb = file_size / (1024 * 1024)
                result = {'event': 'file_done', 'index': idx, 'filename': filename,
                          'status': 'error', 'error': f'File too large ({size_mb:.0f}MB — max 100MB)',
                          'current': idx + 1, 'total': total}
                summaries.append(result)
                os.unlink(tmp_path)
                yield f"data: {json.dumps(result)}\n\n"
                continue

            # --- Read from disk one at a time ---
            # Temp file kept alive until after plaintext supplement
            # (needed for repaired/corrupted files). Deleted at end of
            # each file's processing block regardless of outcome.
            repair_used = False
            try:
                try:
                    with open(tmp_path, 'r', encoding='utf-8', errors='replace') as fh:
                        har_text = fh.read()
                except Exception as e:
                    raise Exception(f'Could not read file: {str(e)}')

                # Option C: clean then del original string
                har_text_clean = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', har_text)
                del har_text

            except Exception as e:
                result = {'event': 'file_done', 'index': idx, 'filename': filename,
                          'status': 'error', 'error': str(e),
                          'current': idx + 1, 'total': total}
                summaries.append(result)
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                yield f"data: {json.dumps(result)}\n\n"
                continue

            # --- Parse ---
            try:
                try:
                    har_data = json.loads(har_text_clean)
                    if 'log' not in har_data or 'entries' not in har_data['log']:
                        raise ValueError('Invalid HAR structure')
                except (json.JSONDecodeError, ValueError):
                    entries, _ = resilient_parse_har(har_text_clean)
                    if not entries:
                        raise Exception('No valid entries found')
                    har_data = {'log': {'version': '1.2', 'creator': {'name': 'Bulk Analyzer'}, 'entries': entries}}
                    repair_used = True

                # Option B: strip response bodies in-place after parse
                strip_response_bodies(har_data)
                del har_text_clean

            except Exception as e:
                result = {'event': 'file_done', 'index': idx, 'filename': filename,
                          'status': 'error', 'error': f'Parse failed: {str(e)}',
                          'current': idx + 1, 'total': total}
                summaries.append(result)
                try:
                    del har_text_clean
                except Exception:
                    pass
                yield f"data: {json.dumps(result)}\n\n"
                continue

            # --- Analyze ---
            try:
                results = analyze_har_simple(har_data)  # dels har_data internally (Option D)

                # Plaintext supplement for repaired/corrupted files.
                # Temp file is still on disk — re-read the cleaned text from it,
                # run the supplement, then delete. This is the fix for the tradeoff
                # noted in the original tempfile implementation.
                if repair_used:
                    try:
                        # Re-read temp file for plaintext supplement.
                        # analyze_har_as_plaintext uses regex pattern matching only —
                        # it does not need pre-cleaned text, so we skip the re.sub
                        # step to avoid holding two full copies of the file in RAM.
                        first_party_domain = results.get('first_party', '')
                        with open(tmp_path, 'r', encoding='utf-8', errors='replace') as fh:
                            supplement_text = fh.read()
                        plaintext_pii = analyze_har_as_plaintext(supplement_text, first_party_domain)
                        del supplement_text
                        if plaintext_pii['pii_count'] > 0:
                            existing = results.get('first_party_pii', {})
                            existing_items = existing.get('pii_items', [])
                            all_items = existing_items + plaintext_pii.get('pii_items', [])
                            unique = {f"{p['type']}:{p['value']}": p for p in all_items}
                            combined = list(unique.values())
                            results['first_party_pii'] = {
                                'pii_items': combined,
                                'pii_count': len(combined),
                                'post_count': existing.get('post_count', 0),
                                'method': 'hybrid' if existing_items else 'plaintext_fallback',
                            }
                    except Exception:
                        pass  # supplement failure is non-fatal
                    finally:
                        # Temp file no longer needed — delete it now
                        try:
                            os.unlink(tmp_path)
                        except Exception:
                            pass
                else:
                    # Not repaired — temp file no longer needed
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass

                total_vendor_pii = sum(v.get('pii_count', 0) for v in results.get('vendors', {}).values())
                first_party_count = results.get('first_party_pii', {}).get('pii_count', 0)
                vendors_with_pii = [v['name'] for v in results.get('vendors', {}).values() if v.get('pii_count', 0) > 0]

                risk = compute_risk_score(results)

                summary = {
                    'event':             'file_done',
                    'index':             idx,
                    'filename':          filename,
                    'status':            'complete',
                    'domain':            results.get('first_party', 'Unknown'),
                    'total_entries':     results.get('total_requests', 0),
                    'vendors_detected':  results.get('vendors_detected', 0),
                    'vendors_with_pii':  vendors_with_pii,
                    'vendor_pii_count':  total_vendor_pii,
                    'first_party_pii_count': first_party_count,
                    'leadid_detected':   results.get('leadid_detected', False),
                    'repair_used':       repair_used,
                    'current':           idx + 1,
                    'total':             total,
                    'score':             risk['score'],
                    'grade':             risk['grade'],
                    'grade_label':       risk['grade_label'],
                    'estimated_damages': risk['estimated_damages'],
                    'top_violations':    risk['top_violations'],
                    'breakdown':         risk['breakdown'],
                }
                summaries.append(summary)

            except Exception as e:
                summary = {'event': 'file_done', 'index': idx, 'filename': filename,
                           'status': 'error', 'error': f'Analysis failed: {str(e)}',
                           'current': idx + 1, 'total': total}
                summaries.append(summary)
            finally:
                try:
                    del results
                except Exception:
                    pass

            yield f"data: {json.dumps(summary)}\n\n"

        # Save bulk session to MongoDB before emitting all_done
        session_id = save_bulk_ranking(summaries)

        # All done
        yield f"data: {json.dumps({'event': 'all_done', 'files_processed': len(summaries), 'summaries': summaries, 'session_id': session_id})}\n\n"

    return app.response_class(
        generate(temp_files),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
        }
    )

# ── Phase 3: Litigation report routes ────────────────────────────────────────

@app.route('/api/generate-report/<analysis_id>', methods=['POST'])
def api_generate_report(analysis_id):
    """
    Generate a litigation package HTML report for a completed analysis.
    Loads the litigation_extract from MongoDB, sends to Claude with the
    saved system prompt and templates, converts Markdown to HTML, stores
    the result, and returns the HTML for download.
    """
    if not CLAUDE_ENABLED:
        return jsonify({'error': 'Claude API not configured — set ANTHROPIC_API_KEY'}), 503

    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503

    # Get model and format from request body
    data        = request.get_json(silent=True) or {}
    model       = data.get('model') or get_admin_settings().get('default_model', DEFAULT_MODEL)
    out_format  = data.get('format', 'html').lower()  # 'html' or 'docx'
    if out_format not in ('html', 'docx'):
        out_format = 'html'

    # Load litigation extract from MongoDB
    doc = _col_single.find_one({'analysis_id': analysis_id}, {'_id': 0})
    if not doc:
        return jsonify({'error': 'Analysis not found'}), 404

    extract = doc.get('litigation_extract')
    if not extract:
        return jsonify({'error': 'No litigation extract found — re-analyze the file to generate one'}), 400

    # Load admin settings (prompt, templates)
    settings        = get_admin_settings()
    system_prompt   = settings.get('system_prompt') or DEFAULT_SYSTEM_PROMPT
    blank_template  = settings.get('blank_template', '')
    sample_ref      = settings.get('sample_reference', '')
    prompt_version  = settings.get('prompt_version', 'v1.0')

    # Build user message
    user_parts = [
        "Below is the structured HAR analysis data. "
        "Produce the complete litigation package following your instructions.\n\n",
        "=" * 60 + "\n",
        "HAR ANALYSIS DATA:\n",
        "=" * 60 + "\n",
        extract,
    ]

    if blank_template:
        user_parts += [
            "\n\n" + "=" * 60 + "\n",
            "OUTPUT STRUCTURE TEMPLATE (follow this section format):\n",
            "=" * 60 + "\n",
            blank_template[:8000],
        ]

    if sample_ref:
        user_parts += [
            "\n\n" + "=" * 60 + "\n",
            "QUALITY REFERENCE EXAMPLE (match this depth and style):\n",
            "=" * 60 + "\n",
            sample_ref[:8000],
        ]

    user_message = ''.join(user_parts)

    # Call Claude API
    try:
        response = _anthropic_client.messages.create(
            model=model,
            max_tokens=12000,
            system=system_prompt,
            messages=[{'role': 'user', 'content': user_message}]
        )

        markdown_output = response.content[0].text if response.content else ''
        input_tokens    = response.usage.input_tokens
        output_tokens   = response.usage.output_tokens

        track_token_usage(input_tokens, output_tokens, model=model)

    except Exception as e:
        print(f"⚠️  Claude API call failed: {e}")
        return jsonify({'error': f'Claude API error: {str(e)}'}), 500

    # Convert Markdown to HTML
    html_content = markdown_to_html(markdown_output)

    # Save report to MongoDB
    report_id = str(uuid.uuid4())
    try:
        _col_reports.insert_one({
            'report_id':      report_id,
            'analysis_id':    analysis_id,
            'created_at':     datetime.now(timezone.utc),
            'model_used':     model,
            'prompt_version': prompt_version,
            'input_tokens':   input_tokens,
            'output_tokens':  output_tokens,
            'status':         'complete',
            'markdown':       markdown_output,   # raw source for both formats
            'html_content':   html_content,
        })
        # Update the single_analyses doc to mark report exists
        _col_single.update_one(
            {'analysis_id': analysis_id},
            {'$set': {'has_report': True, 'report_id': report_id}}
        )
        print(f"✅ Litigation report saved: {report_id}")
    except Exception as e:
        print(f"⚠️  Report save failed (non-fatal): {e}")

    # Return docx as binary download, or JSON with html content
    if out_format == 'docx':
        try:
            from flask import send_file
            docx_buf = markdown_to_docx(markdown_output)
            return send_file(
                docx_buf,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'litigation-package-{report_id[:8]}.docx',
            )
        except Exception as e:
            # Fall back to HTML if docx generation fails
            print(f'⚠️  docx generation failed, falling back to HTML: {e}')

    return jsonify({
        'report_id':     report_id,
        'analysis_id':   analysis_id,
        'model_used':    model,
        'input_tokens':  input_tokens,
        'output_tokens': output_tokens,
        'html':          html_content,
    })


@app.route('/api/reports/count', methods=['GET'])
def api_reports_count():
    """Return total number of litigation reports generated."""
    if not MONGO_ENABLED:
        return jsonify({'count': 0})
    try:
        count = _col_reports.count_documents({})
        return jsonify({'count': count})
    except Exception as e:
        return jsonify({'count': 0, 'error': str(e)})


@app.route('/api/reports/<report_id>', methods=['GET'])
def api_get_report(report_id):
    """
    Retrieve a previously generated litigation report by report_id.
    Returns the stored HTML content for re-download.
    """
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503
    try:
        doc = _col_reports.find_one({'report_id': report_id}, {'_id': 0})
        if not doc:
            return jsonify({'error': 'Report not found'}), 404
        if 'created_at' in doc:
            doc['created_at'] = doc['created_at'].isoformat()
        return jsonify(doc)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/reports/<report_id>/docx', methods=['GET'])
def api_get_report_docx(report_id):
    """
    Re-download a previously generated report as a .docx file.
    Converts the stored raw markdown to docx on the fly.
    """
    if not MONGO_ENABLED:
        return jsonify({'error': 'Database not available'}), 503
    try:
        from flask import send_file
        doc = _col_reports.find_one({'report_id': report_id}, {'_id': 0, 'markdown': 1})
        if not doc:
            return jsonify({'error': 'Report not found'}), 404
        md = doc.get('markdown', '')
        if not md:
            return jsonify({'error': 'No markdown source stored — regenerate the report'}), 400
        docx_buf = markdown_to_docx(md)
        return send_file(
            docx_buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'litigation-package-{report_id[:8]}.docx',
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/usage/summary', methods=['GET'])
def api_admin_usage_summary():
    """
    Return token usage summary for the last 30 days and all time.
    Groups by model and computes estimated cost using per-model rates.
    """
    # Per-model pricing ($ per million tokens)
    MODEL_RATES = {
        'claude-haiku-4-5-20251001': {'input': 1.00,  'output': 5.00},
        'claude-haiku-4-5':          {'input': 1.00,  'output': 5.00},
        'claude-sonnet-4-6':         {'input': 3.00,  'output': 15.00},
        'claude-opus-4-6':           {'input': 5.00,  'output': 25.00},
        'unknown':                   {'input': 3.00,  'output': 15.00},  # assume Sonnet if unknown
    }

    def compute_cost(model, inp, out):
        rates = MODEL_RATES.get(model, MODEL_RATES['unknown'])
        return (inp / 1_000_000 * rates['input']) + (out / 1_000_000 * rates['output'])

    # All-time totals from running-total doc
    all_time = get_token_usage()
    all_time_inp  = all_time.get('total_input_tokens', 0)
    all_time_out  = all_time.get('total_output_tokens', 0)
    all_time_calls = all_time.get('total_calls', 0)

    if not MONGO_ENABLED:
        return jsonify({
            'db_enabled':    False,
            'all_time':      {'input': all_time_inp, 'output': all_time_out,
                              'calls': all_time_calls, 'cost': 0},
            'last_30_days':  {'input': 0, 'output': 0, 'calls': 0,
                              'cost': 0, 'by_model': []},
        })

    try:
        cutoff = datetime.now(timezone.utc) - timedelta(days=30)

        # Aggregate last-30-day calls grouped by model
        pipeline = [
            {'$match': {'timestamp': {'$gte': cutoff}}},
            {'$group': {
                '_id':          '$model',
                'input_tokens': {'$sum': '$input_tokens'},
                'output_tokens':{'$sum': '$output_tokens'},
                'calls':        {'$sum': 1},
            }},
            {'$sort': {'input_tokens': -1}},
        ]
        rows = list(_col_token_calls.aggregate(pipeline))

        by_model = []
        total_30_inp = total_30_out = total_30_calls = 0
        total_30_cost = 0.0

        for row in rows:
            model  = row['_id'] or 'unknown'
            inp    = row['input_tokens']
            out    = row['output_tokens']
            calls  = row['calls']
            cost   = compute_cost(model, inp, out)
            total_30_inp   += inp
            total_30_out   += out
            total_30_calls += calls
            total_30_cost  += cost
            by_model.append({
                'model':         model,
                'input_tokens':  inp,
                'output_tokens': out,
                'calls':         calls,
                'cost':          round(cost, 4),
            })

        # All-time estimated cost (best-effort using same rates)
        # Since we don't have per-model breakdown all-time, use 30-day mix
        # or fall back to Sonnet rate for unknown distribution
        all_time_cost = round(compute_cost('unknown', all_time_inp, all_time_out), 4)

        return jsonify({
            'db_enabled':   True,
            'all_time': {
                'input':  all_time_inp,
                'output': all_time_out,
                'calls':  all_time_calls,
                'cost':   all_time_cost,
            },
            'last_30_days': {
                'input':    total_30_inp,
                'output':   total_30_out,
                'calls':    total_30_calls,
                'cost':     round(total_30_cost, 4),
                'by_model': by_model,
            },
        })
    except Exception as e:
        print(f'⚠️  Usage summary failed: {e}')
        return jsonify({'error': str(e)}), 500


# ── Claude / Anthropic helpers ───────────────────────────────────────────────

def track_token_usage(input_tokens, output_tokens, model=None):
    """
    Accumulate token usage in MongoDB.
    Writes to two collections:
      - token_usage: single running-total doc (all-time)
      - token_calls: one doc per API call with timestamp + model (for 30-day breakdown)
    Non-fatal if DB unavailable.
    """
    if not MONGO_ENABLED:
        return
    now = datetime.now(timezone.utc)
    try:
        # Running totals
        _col_token_use.update_one(
            {'_id': 'global'},
            {
                '$inc': {
                    'total_input_tokens':  input_tokens,
                    'total_output_tokens': output_tokens,
                    'total_calls':         1,
                },
                '$set': {'last_updated': now},
            },
            upsert=True
        )
        # Per-call record
        _col_token_calls.insert_one({
            'timestamp':     now,
            'model':         model or 'unknown',
            'input_tokens':  input_tokens,
            'output_tokens': output_tokens,
        })
    except Exception as e:
        print(f'⚠️  Token tracking failed: {e}')


def get_token_usage():
    """Return cumulative token usage from MongoDB, or zeros if unavailable."""
    if not MONGO_ENABLED:
        return {'total_input_tokens': 0, 'total_output_tokens': 0, 'total_calls': 0}
    try:
        doc = _col_token_use.find_one({'_id': 'global'}) or {}
        return {
            'total_input_tokens':  doc.get('total_input_tokens',  0),
            'total_output_tokens': doc.get('total_output_tokens', 0),
            'total_calls':         doc.get('total_calls',         0),
        }
    except Exception:
        return {'total_input_tokens': 0, 'total_output_tokens': 0, 'total_calls': 0}


@app.route('/api/claude/status', methods=['GET'])
def api_claude_status():
    """
    Return Claude API status and token usage.

    ?test=1  — actually sends a live ping to the API (costs tokens).
               Only called when the user clicks 'Test Connection'.
    No param  — returns configuration status and token counts only,
               no API call made. Safe to call freely.
    """
    api_key_set = bool(os.environ.get('ANTHROPIC_API_KEY', ''))
    run_test    = request.args.get('test') == '1'

    if not CLAUDE_ENABLED or not run_test:
        return jsonify({
            'connected':   CLAUDE_ENABLED,
            'api_key_set': api_key_set,
            'live_test':   False,
            'error':       None if CLAUDE_ENABLED else (
                           'ANTHROPIC_API_KEY not set' if not api_key_set
                           else 'anthropic package not installed'),
            'usage':       get_token_usage(),
        })

    # ?test=1 — send a live ping (user explicitly requested)
    try:
        response = _anthropic_client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=5,
            messages=[{'role': 'user', 'content': 'Reply with one word: OK'}]
        )
        reply         = response.content[0].text.strip() if response.content else ''
        input_tokens  = response.usage.input_tokens
        output_tokens = response.usage.output_tokens

        track_token_usage(input_tokens, output_tokens, model='claude-haiku-4-5-20251001')

        return jsonify({
            'connected':    True,
            'api_key_set':  True,
            'live_test':    True,
            'model_tested': 'claude-haiku-4-5-20251001',
            'reply':        reply,
            'test_tokens':  {'input': input_tokens, 'output': output_tokens},
            'error':        None,
            'usage':        get_token_usage(),
        })

    except Exception as e:
        return jsonify({
            'connected':   False,
            'api_key_set': api_key_set,
            'live_test':   True,
            'error':       str(e),
            'usage':       get_token_usage(),
        })


if __name__ == '__main__':
    # This file should be run via app.py
    # For direct execution during development:
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)

